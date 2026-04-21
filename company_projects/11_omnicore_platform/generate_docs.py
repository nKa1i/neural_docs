"""
Script to generate the large DOCX and PDF test files for 11_omnicore_platform.
Run once: python generate_docs.py
"""
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── DOCX ─────────────────────────────────────────────────────────────────────

def make_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    doc.add_heading("OmniCore Platform — Product Requirements Document (PRD)", 0)
    doc.add_paragraph(
        "Document Type: Product Requirements Document | Version: 4.1 | "
        "Status: DRAFT — NOT FINAL | Owner: Maria Svetlova (CPO) | "
        "Last Updated: 2025-02-25 | Review Due: 2025-03-10"
    )

    # Executive Summary
    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "OmniCore Platform is a next-generation cloud ERP designed to serve enterprise "
        "manufacturers, distributors, and logistics companies across the CIS and MENA regions. "
        "The platform consolidates 17 business modules into a single AI-powered system with "
        "real-time analytics, predictive inventory management, and full regulatory compliance "
        "for Russia (152-FZ), Kazakhstan (PDP Law 2023), and the European Union (GDPR)."
    )
    doc.add_paragraph(
        "This PRD defines the full functional requirements for OmniCore Platform v1.0 (GA release). "
        "It supersedes all previous requirement documents (PRD v1.x through v3.x). "
        "Any conflicts with architecture.md should be resolved in favour of this document "
        "unless explicitly noted."
    )

    # Business Goals
    doc.add_heading("2. Business Goals and Success Metrics", 1)
    doc.add_heading("2.1 Strategic Goals", 2)
    goals = [
        ("G1", "Capture 2% of the CIS ERP market within 3 years of GA release."),
        ("G2", "Achieve ARR of $18,000,000 by end of Year 2 (2027)."),
        ("G3", "Maintain NPS ≥ 50 across all customer tiers."),
        ("G4", "Process 1 billion financial transactions per year by 2028."),
        ("G5", "Zero critical security incidents (CVSS ≥ 9.0) in production."),
    ]
    for code, text in goals:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{code}: ").bold = True
        p.add_run(text)

    doc.add_heading("2.2 KPIs", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Target (Year 1)"
    hdr[2].text = "Target (Year 2)"
    hdr[3].text = "Measurement Method"
    kpi_data = [
        ("Monthly Active Companies", "500", "1,500", "Billing system"),
        ("API Uptime", "99.95%", "99.99%", "Synthetic monitoring"),
        ("P99 API Latency", "< 200ms", "< 150ms", "Prometheus/Grafana"),
        ("Time to Onboard New Client", "≤ 3 weeks", "≤ 2 weeks", "CRM tracking"),
        ("Support Ticket CSAT", "≥ 4.2 / 5", "≥ 4.5 / 5", "Intercom surveys"),
        ("Data Export Availability", "< 5 min", "< 2 min", "Job queue monitoring"),
        ("Mobile App Crash Rate", "< 0.1%", "< 0.05%", "Firebase Crashlytics"),
    ]
    for row_data in kpi_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    # Pricing Section (CONFLICT WITH OTHER DOCS)
    doc.add_heading("3. Pricing Model (CPO Version — DRAFT)", 1)
    doc.add_paragraph(
        "NOTE: This pricing model is proposed by the CPO and has NOT been approved by the CFO. "
        "The CFO has a different pricing model in budget_draft.txt. Final decision pending."
    )
    doc.add_heading("3.1 Subscription Tiers", 2)
    pricing_table = doc.add_table(rows=1, cols=5)
    pricing_table.style = "Table Grid"
    ph = pricing_table.rows[0].cells
    for i, h in enumerate(["Tier", "Price/Month", "Users", "Modules", "SLA"]):
        ph[i].text = h
    pricing_rows = [
        ("Starter", "$1,990", "Up to 50", "5 core modules", "99.9%"),
        ("Business", "$5,990", "Up to 200", "12 modules", "99.95%"),
        ("Enterprise", "$14,990", "Unlimited", "All 17 modules", "99.99%"),
        ("AI Pack (add-on)", "$990", "N/A", "AI Engine + BI Premium", "Included"),
    ]
    for pr in pricing_rows:
        row = pricing_table.add_row().cells
        for i, val in enumerate(pr):
            row[i].text = val

    # Module Requirements (long section)
    doc.add_heading("4. Detailed Module Requirements", 1)

    modules = [
        ("4.1 Finance Core Module", [
            "FR-FIN-001: System SHALL support double-entry bookkeeping with automatic balance validation.",
            "FR-FIN-002: System SHALL support 45 currencies with real-time exchange rate integration (ECB, NBK, CBR).",
            "FR-FIN-003: System SHALL generate IFRS-compliant consolidated financial statements for up to 30 subsidiaries.",
            "FR-FIN-004: System SHALL perform automatic bank reconciliation with tolerance of ±$0.01.",
            "FR-FIN-005: System SHALL support 5-year budget planning with quarterly reforecast capability.",
            "FR-FIN-006: System SHALL detect financial anomalies using ML (Isolation Forest) with < 1% false positive rate.",
            "FR-FIN-007: System SHALL generate electronic invoices compliant with Kazakhstan e-invoice standard (ЭСФ) and Russian EDI (EDIFACT/X12).",
            "FR-FIN-008: System SHALL archive financial documents for 10 years per legal requirements.",
            "FR-FIN-009: System SHALL integrate with minimum 12 banking partners via Open Banking API (ISO 20022).",
            "FR-FIN-010: System SHALL provide 30-day cash flow forecasting with 85% accuracy (MAPE ≤ 15%).",
        ]),
        ("4.2 Manufacturing Core Module", [
            "FR-MFG-001: System SHALL support MRP II planning (Material Requirements Planning level 2).",
            "FR-MFG-002: System SHALL manage Bills of Materials (BOM) with up to 15 nesting levels.",
            "FR-MFG-003: System SHALL track production orders from creation to completion with full audit trail.",
            "FR-MFG-004: System SHALL integrate with SCADA systems via OPC-UA protocol.",
            "FR-MFG-005: System SHALL maintain digital twins for production lines with real-time sensor data.",
            "FR-MFG-006: System SHALL enforce quality control checkpoints per ISO 9001:2015.",
            "FR-MFG-007: System SHALL calculate OEE (Overall Equipment Effectiveness) in real-time.",
            "FR-MFG-008: System SHALL support sub-contracting operations with material tracking.",
        ]),
        ("4.3 Warehouse Management System (WMS)", [
            "FR-WMS-001: System SHALL support bin-level addressing with 3D warehouse visualization.",
            "FR-WMS-002: System SHALL process RFID scans at ≥ 10,000 scans/second per reader cluster.",
            "FR-WMS-003: System SHALL support automated drone-based inventory counting.",
            "FR-WMS-004: System SHALL implement FIFO, LIFO, and FEFO lot management strategies.",
            "FR-WMS-005: System SHALL perform ABC analysis on warehouse SKUs.",
            "FR-WMS-006: System SHALL optimise pick routes using S-shape or return routing algorithms.",
            "FR-WMS-007: System SHALL support pick waves of up to 150 orders simultaneously.",
            "FR-WMS-008: System SHALL integrate with conveyor systems via REST API.",
            "FR-WMS-009: System SHALL provide slotting optimisation recommendations based on velocity.",
            "FR-WMS-010: System SHALL generate shipping labels in ZPL, PDF, and GS1-128 formats.",
        ]),
        ("4.4 Supply Chain Management", [
            "FR-SCM-001: System SHALL optimise delivery routes using VRP (Vehicle Routing Problem) algorithm.",
            "FR-SCM-002: System SHALL manage supplier relationships with scorecards and risk ratings.",
            "FR-SCM-003: System SHALL support electronic tendering and reverse auction.",
            "FR-SCM-004: System SHALL forecast demand using LSTM and Prophet ML models.",
            "FR-SCM-005: System SHALL provide GPS/IoT-based shipment track & trace.",
            "FR-SCM-006: System SHALL integrate with SAP TM and Oracle TMS.",
            "FR-SCM-007: System SHALL visualise supply chain as a graph (Neo4j) for risk analysis.",
            "FR-SCM-008: System SHALL calculate supplier lead times and safety stock automatically.",
        ]),
        ("4.5 HR & Payroll Module", [
            "FR-HR-001: System SHALL calculate payroll for employees in 50+ tax jurisdictions.",
            "FR-HR-002: System SHALL integrate with Learning Management Systems (LMS) via xAPI/SCORM.",
            "FR-HR-003: System SHALL support OKR and KPI tracking with quarterly review cycles.",
            "FR-HR-004: System SHALL provide AI-powered resume screening with 90% precision on shortlisting.",
            "FR-HR-005: System SHALL maintain compliance with local labour law for KZ, RU, EU jurisdictions.",
        ]),
        ("4.6 CRM Module", [
            "FR-CRM-001: System SHALL maintain 360° customer profiles aggregating data from all modules.",
            "FR-CRM-002: System SHALL predict customer churn with AUC-ROC ≥ 0.85.",
            "FR-CRM-003: System SHALL support CPQ (Configure, Price, Quote) with dynamic pricing rules.",
            "FR-CRM-004: System SHALL segment customers using RFM analysis and ML clustering.",
            "FR-CRM-005: System SHALL automate marketing campaigns with A/B testing.",
        ]),
        ("4.7 BI & Analytics Module", [
            "FR-BI-001: System SHALL provide OLAP cubes with drag-and-drop query builder.",
            "FR-BI-002: System SHALL update real-time dashboards within 30 seconds of data change.",
            "FR-BI-003: System SHALL support natural language queries (NLQ) via GPT-4o integration.",
            "FR-BI-004: System SHALL export reports in Excel (.xlsx), PDF, Power BI (.pbix), and CSV formats.",
            "FR-BI-005: System SHALL provide anomaly alerts with root cause explanation.",
            "FR-BI-006: System SHALL support embedded analytics (iframe embed with SSO).",
        ]),
    ]

    for title, reqs in modules:
        doc.add_heading(title, 2)
        for req in reqs:
            doc.add_paragraph(req, style="List Bullet")

    # Non-functional requirements
    doc.add_heading("5. Non-Functional Requirements", 1)
    nfr_sections = [
        ("5.1 Performance", [
            "NFR-PERF-001: REST API P99 response time MUST be ≤ 200ms under normal load (≤ 50,000 RPS).",
            "NFR-PERF-002: ClickHouse analytical queries MUST return within 5 seconds for 95% of queries.",
            "NFR-PERF-003: Kafka event throughput MUST support ≥ 500,000 events/second.",
            "NFR-PERF-004: System MUST support 100,000 concurrent users across all tenants.",
            "NFR-PERF-005: File export jobs (Excel, PDF) MUST complete within 60 seconds for reports up to 100,000 rows.",
            "NFR-PERF-006: Database query plans MUST be reviewed for any query exceeding 100ms (flagged by PgBouncer slow log).",
        ]),
        ("5.2 Security", [
            "NFR-SEC-001: All data in transit MUST be encrypted with TLS 1.3 minimum.",
            "NFR-SEC-002: All data at rest MUST be encrypted with AES-256 (AWS KMS).",
            "NFR-SEC-003: PII fields (SSN, bank accounts, passport numbers) MUST be encrypted at field level.",
            "NFR-SEC-004: Authentication MUST support OAuth 2.0 + OIDC, SAML 2.0, and LDAP/Active Directory.",
            "NFR-SEC-005: All database credentials MUST be rotated every 24 hours via HashiCorp Vault.",
            "NFR-SEC-006: System MUST pass annual penetration testing with no Critical or High findings unresolved.",
            "NFR-SEC-007: All user actions MUST be recorded in an immutable audit log.",
            "NFR-SEC-008: System MUST comply with OWASP Top 10 (latest edition).",
        ]),
        ("5.3 Reliability & Availability", [
            "NFR-REL-001: Enterprise tier SLA MUST be 99.99% monthly uptime.",
            "NFR-REL-002: Starter/Business tier SLA MUST be 99.95% monthly uptime.",
            "NFR-REL-003: RTO (Recovery Time Objective) for Enterprise MUST be ≤ 15 minutes.",
            "NFR-REL-004: RPO (Recovery Point Objective) for Enterprise MUST be ≤ 5 minutes.",
            "NFR-REL-005: Planned maintenance windows MUST be scheduled with 7-day notice and completed within 2-hour windows.",
            "NFR-REL-006: Kafka consumer lag MUST NOT exceed 100,000 messages for > 10 consecutive minutes.",
        ]),
        ("5.4 Scalability", [
            "NFR-SCAL-001: System MUST auto-scale horizontally (Kubernetes HPA) based on CPU > 70% and memory > 80%.",
            "NFR-SCAL-002: Database connection pool MUST handle up to 500 concurrent connections per service.",
            "NFR-SCAL-003: System MUST support up to 10,000 tenants with data isolation.",
            "NFR-SCAL-004: Single tenant MUST be scalable to 10,000 concurrent users.",
        ]),
        ("5.5 Usability", [
            "NFR-USE-001: Web interface MUST load initial page in < 2 seconds on a 10 Mbps connection.",
            "NFR-USE-002: Mobile app (iOS/Android) MUST support offline mode for field operations (WMS, inventory).",
            "NFR-USE-003: Interface MUST be localised in Russian, Kazakh, and English.",
            "NFR-USE-004: All forms MUST be accessible per WCAG 2.1 Level AA.",
            "NFR-USE-005: User onboarding wizard MUST allow a new user to complete first task within 15 minutes.",
        ]),
    ]
    for title, reqs in nfr_sections:
        doc.add_heading(title, 2)
        for req in reqs:
            doc.add_paragraph(req, style="List Bullet")

    # Timeline conflict section
    doc.add_heading("6. Project Timeline", 1)
    doc.add_paragraph(
        "CONFLICT NOTE: Multiple timeline versions exist across documents. "
        "This PRD reflects the CPO's proposed timeline. See architecture.md (ADR-001) for the full conflict."
    )
    timeline_table = doc.add_table(rows=1, cols=4)
    timeline_table.style = "Table Grid"
    th = timeline_table.rows[0].cells
    for i, h in enumerate(["Phase", "Duration", "End Date (CPO)", "Deliverables"]):
        th[i].text = h
    timeline_data = [
        ("Phase 0: Discovery", "3 months", "March 2025", "Architecture ADRs, team onboarding"),
        ("Phase 1: MVP (5 modules)", "7 months", "October 2025", "Finance, WMS, HR, CRM, BI"),
        ("Phase 2: Full Platform", "7 months", "May 2026", "All 17 modules"),
        ("Beta Testing", "2 months", "July 2026", "50 pilot customers"),
        ("GA Release", "—", "August 2026", "Public launch"),
    ]
    for row_data in timeline_data:
        row = timeline_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph(
        "Total project duration (CPO version): 19 months from January 2025 to August 2026."
    )

    # Appendix
    doc.add_heading("Appendix A: Glossary", 1)
    glossary = [
        ("ERP", "Enterprise Resource Planning — integrated management of core business processes."),
        ("WMS", "Warehouse Management System — software for managing warehouse operations."),
        ("MRP II", "Manufacturing Resource Planning level 2 — production planning methodology."),
        ("CQRS", "Command Query Responsibility Segregation — architectural pattern separating reads/writes."),
        ("VRP", "Vehicle Routing Problem — combinatorial optimisation for delivery route planning."),
        ("OEE", "Overall Equipment Effectiveness — manufacturing productivity metric."),
        ("FIFO/LIFO/FEFO", "First In First Out / Last In First Out / First Expiry First Out — lot strategies."),
        ("BOM", "Bill of Materials — list of raw materials and components needed to manufacture a product."),
        ("CPQ", "Configure, Price, Quote — sales automation for complex product pricing."),
        ("ARR", "Annual Recurring Revenue — annualised subscription revenue metric."),
        ("NPS", "Net Promoter Score — customer loyalty metric (-100 to +100)."),
        ("RTO", "Recovery Time Objective — maximum acceptable downtime after an incident."),
        ("RPO", "Recovery Point Objective — maximum acceptable data loss measured in time."),
        ("SLA", "Service Level Agreement — contractual uptime and performance guarantees."),
        ("GDPR", "General Data Protection Regulation — EU data privacy regulation."),
        ("152-FZ", "Russian Federal Law on Personal Data."),
        ("EDI", "Electronic Data Interchange — standard format for business document exchange."),
        ("OIDC", "OpenID Connect — authentication layer on top of OAuth 2.0."),
        ("mTLS", "Mutual TLS — both client and server authenticate via certificates."),
        ("OLAP", "Online Analytical Processing — database approach for complex queries on large datasets."),
        ("OLTP", "Online Transaction Processing — database approach for fast transactional operations."),
        ("BPMN", "Business Process Model and Notation — standard for business process diagrams."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)

    doc.add_heading("Appendix B: Open Issues", 1)
    issues = [
        "ISSUE-001: Pricing model not finalized — 3 competing proposals (CPO: $1,990/$5,990/$14,990; CFO: $2,500/$7,500/$18,000; Sales: $2,200/$6,500/$16,500).",
        "ISSUE-002: Timeline conflict — CPO: GA August 2026 (19 months); CTO: GA August 2026 (20 months); Architect: GA November 2026 (22 months).",
        "ISSUE-003: Total budget not agreed — CEO: $6,800,000; CFO: $4,200,000.",
        "ISSUE-004: Team size disputed — CEO wants 47 people; CFO approved 32.",
        "ISSUE-005: Frontend architecture not decided — single Next.js monorepo vs micro-frontend with Module Federation.",
        "ISSUE-006: AI provider cost model not approved — GPT-4o API vs self-hosted Llama 3 70B.",
        "ISSUE-007: Kazakhtelecom Cloud deployment timeline — Q3 2026 target but procurement not started.",
        "ISSUE-008: 1C migration partnership — term sheet not yet signed.",
    ]
    for issue in issues:
        doc.add_paragraph(issue, style="List Bullet")

    out_path = os.path.join(OUT_DIR, "product_requirements.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


# ─── PDF ──────────────────────────────────────────────────────────────────────

def make_pdf():
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(100, 100, 100)
            self.cell(0, 8, "OmniCore Platform - Technical Specification v2.1 | CONFIDENTIAL", align="C")
            self.ln(4)
            self.set_draw_color(200, 200, 200)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(3)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Page {self.page_no()} | OmniCore Technologies LLP | Confidential", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(15, 20, 15)
    pdf.add_page()

    def h1(text):
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(30, 70, 140)
        pdf.ln(4)
        pdf.cell(0, 10, text, ln=True)
        pdf.set_draw_color(30, 70, 140)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(3)
        pdf.set_text_color(0, 0, 0)

    def h2(text):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(50, 90, 160)
        pdf.ln(3)
        pdf.cell(0, 8, text, ln=True)
        pdf.set_text_color(0, 0, 0)

    def h3(text):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(0, 7, text, ln=True)
        pdf.set_text_color(0, 0, 0)

    def body(text):
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5.5, text)
        pdf.ln(1)

    def bullet(text, indent=5):
        pdf.set_font("Helvetica", "", 10)
        pdf.set_x(15 + indent)
        pdf.cell(5, 5.5, chr(149), ln=False)
        pdf.multi_cell(0, 5.5, text)

    def table_row(cells, widths, bold=False):
        style = "B" if bold else ""
        pdf.set_font("Helvetica", style, 9)
        x = pdf.get_x()
        y = pdf.get_y()
        max_h = 6
        for cell, w in zip(cells, widths):
            pdf.set_xy(x, y)
            pdf.cell(w, max_h, str(cell), border=1)
            x += w
        pdf.ln(max_h)

    # Title page content
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(30, 70, 140)
    pdf.cell(0, 15, "OmniCore Platform", ln=True, align="C")
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Technical Specification Document", ln=True, align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 7, "Version 2.1 | February 2025 | DRAFT - CONFIDENTIAL", ln=True, align="C")
    pdf.ln(5)
    pdf.set_font("Helvetica", "", 10)
    meta_lines = [
        ("Prepared by:", "Denis Krasnov, Chief Architect; Kirill Osipov, CTO"),
        ("Reviewed by:", "Artur Zhandarov (CEO), Maria Svetlova (CPO)"),
        ("Document status:", "DRAFT - pending board approval"),
        ("Classification:", "CONFIDENTIAL - Internal Use Only"),
        ("Distribution:", "Executive team, Tech leads, Investors (on request)"),
    ]
    pdf.set_text_color(0, 0, 0)
    for label, value in meta_lines:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(50, 6, label)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, value, ln=True)
    pdf.ln(5)

    # Section 1: Overview
    h1("1. System Overview")
    body(
        "OmniCore Platform is a multi-tenant, cloud-native Enterprise Resource Planning (ERP) "
        "system targeting the manufacturing and distribution sectors in Kazakhstan, Russia, and "
        "the broader CIS region. The platform is built on a microservices architecture with 47 "
        "independent services deployed on Kubernetes across 3 cloud regions."
    )
    body(
        "Core differentiation: (1) Native AI layer - ML models are first-class citizens, not add-ons; "
        "(2) Time-to-value of 3 weeks versus 6-18 months for SAP/Oracle; "
        "(3) Low-code BPMN 2.0 process configurator; (4) Full data sovereignty (RU/KZ/EU)."
    )

    h2("1.1 System Boundaries")
    body("The platform boundary includes the following 17 business modules:")
    module_list = [
        "Finance Core (GL, AP, AR, reconciliation, multi-currency)",
        "Manufacturing Core (MRP II, production orders, BOM, digital twins)",
        "Warehouse Management System (WMS, RFID, drones, slotting)",
        "Supply Chain Management (VRP routing, SRM, demand forecasting)",
        "HR & Payroll (50+ tax jurisdictions, OKR tracking, ATS)",
        "CRM (360 profile, churn prediction, CPQ, segmentation)",
        "BI & Analytics (OLAP, NLQ via GPT-4o, real-time dashboards)",
        "Document Management (DMS, OCR, versioning)",
        "Project Management (Gantt, milestones, resource planning)",
        "Asset Management (fixed assets, depreciation, CAPEX tracking)",
        "Procurement (PO management, vendor portal, e-tendering)",
        "Quality Control (ISO 9001:2015, inspection plans, NCR tracking)",
        "Maintenance Management (PPM, work orders, spare parts)",
        "Transport Management (TMS integration, fleet management)",
        "Customs & Compliance (HS codes, customs declarations, KZ/RU forms)",
        "AI Engine (LLM orchestration, embeddings, document intelligence)",
        "Notification Hub (email, SMS, push, Slack, Teams)",
    ]
    for m in module_list:
        bullet(m)

    # Section 2: Architecture
    h1("2. Architecture Overview")
    h2("2.1 Technology Stack Summary")
    tech_rows = [
        ("Layer", "Technology", "Version", "Notes"),
        ("API Gateway", "Kong", "3.6", "Rate limiting, JWT validation, routing"),
        ("Service Mesh", "Istio", "1.21", "mTLS, traffic management, observability"),
        ("Container Platform", "Kubernetes", "1.30", "EKS (AWS), managed"),
        ("Backend Primary", "Go (Fiber v2)", "1.22 / v2.52", "All 40 Go microservices"),
        ("Backend Secondary", "Python (FastAPI)", "3.12 / 0.111", "ML workers, analytics"),
        ("Backend Legacy", "Java (Spring Boot)", "21 / 3.3", "Legacy ERP adapters"),
        ("Frontend", "React + Next.js", "18.3 / 14.2", "TypeScript 5.4, SSR+CSR"),
        ("OLTP Database", "PostgreSQL", "16.2", "Patroni HA, 3 nodes"),
        ("OLAP Database", "ClickHouse", "24.3", "3 shards x 2 replicas"),
        ("Cache", "Redis", "7.2", "Sentinel, 3 nodes"),
        ("Event Streaming", "Apache Kafka", "3.7", "5 brokers, replication=3"),
        ("Search", "Elasticsearch", "8.13", "Full-text, autocomplete"),
        ("ML Serving", "Triton Inference", "24.05", "GPU inference, ONNX"),
        ("IaC", "Terraform", "1.8", "All infra as code"),
        ("GitOps", "ArgoCD", "2.11", "Declarative deployments"),
        ("Secrets", "HashiCorp Vault", "1.16", "Dynamic secrets, rotation"),
        ("Monitoring", "Prometheus/Grafana", "2.52 / 11.0", "Metrics, dashboards, SLOs"),
        ("Tracing", "Jaeger + OTEL", "1.57", "Distributed tracing"),
        ("Logging", "Loki + Promtail", "3.0", "Structured logs, 30-day retention"),
    ]
    col_w = [38, 36, 22, 74]
    table_row([r[0] for r in tech_rows[:1]], col_w, bold=True)
    for row in tech_rows[1:]:
        table_row(list(row), col_w)
    pdf.ln(3)

    h2("2.2 Service Count and Scaling")
    body(
        "The platform comprises 47 microservices (CONFLICT: CTO proposes to consolidate to 32 - "
        "decision pending as of Feb 2025, see ADR-002 in architecture.md). Each service is "
        "independently deployable with its own CI/CD pipeline, database schema, and scaling policy."
    )
    body(
        "Scaling policy: Kubernetes Horizontal Pod Autoscaler (HPA) is configured for all services "
        "with CPU target 70% and memory target 80%. Minimum replicas: 2 (all environments except dev). "
        "Maximum replicas: 20 (Go services), 10 (Python/Java services), 6 (ML inference)."
    )

    h2("2.3 Data Flow Architecture")
    body("Primary request path (synchronous, user-facing):")
    data_flow = [
        "1. Client request arrives at CloudFront CDN (TLS termination, static asset cache).",
        "2. WAF inspection (AWS WAF + ModSecurity rules, OWASP Top 10 protection).",
        "3. Kong API Gateway: JWT validation, rate limit check (1K/10K/unlimited by tier), routing.",
        "4. Auth Service: token introspection, RBAC permission check.",
        "5. Target microservice processes the request (business logic, DB read/write).",
        "6. Response returned with X-Request-ID, X-Trace-ID headers.",
        "7. Async: Kafka event published for downstream consumers (BI, audit, notifications).",
    ]
    for step in data_flow:
        bullet(step)

    # Section 3: Security
    h1("3. Security Architecture")
    h2("3.1 Authentication & Authorisation")
    body(
        "Identity provider: Keycloak 24 (self-hosted, HA, 3-node cluster). "
        "Supports: OAuth 2.0 + OIDC (primary), SAML 2.0 (enterprise SSO), LDAP/Active Directory (on-premise). "
        "Token TTL: access token 15 minutes, refresh token 30 days. "
        "MFA: TOTP (Google Authenticator / Authy) and WebAuthn/FIDO2 hardware keys."
    )
    body("RBAC model: 5 built-in roles - super_admin, tenant_admin, module_admin, operator, viewer. "
         "Custom roles supported via permission matrix. Row-level security enforced via tenant_id "
         "predicate on all database queries.")

    h2("3.2 Data Protection")
    sec_items = [
        "Encryption at rest: AES-256-GCM via AWS KMS (automatic key rotation every 365 days).",
        "Encryption in transit: TLS 1.3 minimum on all connections (clients, service-to-service, DB).",
        "PII field encryption: SSN, bank account, passport numbers encrypted at field level (AES-GCM, per-tenant key).",
        "Data masking: PII masked in non-production environments (dev, staging) via automated tooling.",
        "Secrets management: HashiCorp Vault - database credentials rotated every 24 hours, no static passwords.",
        "Vulnerability scanning: Trivy (containers) + Snyk (dependencies) in every CI pipeline run.",
        "SAST: SonarQube - quality gate blocks merge if BLOCKER or CRITICAL issues found.",
        "Annual penetration testing: external firm, CVSS-based severity classification.",
    ]
    for item in sec_items:
        bullet(item)

    # Section 4: Performance
    h1("4. Performance Specifications")
    h2("4.1 Load Targets")
    load_rows = [
        ("Metric", "Value", "Measurement Condition"),
        ("Total concurrent users", "100,000", "Across all tenants, all tiers"),
        ("API RPS (steady state)", "50,000 RPS", "Normal business hours"),
        ("API RPS (peak, 3x multiplier)", "150,000 RPS", "Month-end close, peak season"),
        ("P99 API response time", "<= 200ms", "GET /v1/* endpoints under steady load"),
        ("P99 POST (write) response time", "<= 500ms", "Write operations with DB commit"),
        ("Kafka events/second", "500,000 ev/s", "All topics combined"),
        ("ClickHouse queries/second", "500 concurrent", "OLAP analytical queries"),
        ("Elasticsearch queries/second", "5,000 QPS", "Full-text search"),
        ("File export (100K rows Excel)", "<= 60 seconds", "Background job"),
        ("Database P99 query latency", "<= 20ms", "Indexed queries, OLTP only"),
    ]
    col_w2 = [62, 40, 68]
    table_row([r[0] for r in load_rows[:1]], col_w2, bold=True)
    for row in load_rows[1:]:
        table_row(list(row), col_w2)
    pdf.ln(3)

    h2("4.2 Caching Architecture")
    body(
        "Multi-layer caching strategy to achieve P99 <= 200ms: "
        "(1) CloudFront CDN - static assets (1h TTL), public read endpoints (60s TTL); "
        "(2) Kong API Gateway cache (Redis) - read-heavy endpoints, 60s TTL; "
        "(3) Application Redis cache - user sessions (24h), RBAC decisions (5min), config (5min), dashboard KPIs (30s); "
        "(4) PgBouncer connection pool - eliminates per-request connection overhead."
    )

    # Section 5: Infrastructure
    h1("5. Infrastructure & Cloud")
    h2("5.1 Multi-Region Deployment")
    body(
        "Primary region: AWS us-east-1 (Virginia) - full stack, primary write cluster. "
        "Secondary region: AWS eu-central-1 (Frankfurt) - EU GDPR boundary, full stack, EU data residency. "
        "Tertiary region: Yandex Cloud ru-central1 (Moscow) - 152-FZ compliance, air-gapped from AWS, "
        "RU customer data never leaves this region. "
        "Planned Q3 2026: Kazakhtelecom Cloud (Almaty) - KZ government compliance."
    )

    h2("5.2 Kubernetes Node Pools (Production)")
    node_rows = [
        ("Pool", "Instance", "Count", "Purpose"),
        ("system-pool", "m6i.xlarge (4 vCPU, 16 GB)", "3", "Kong, Istio, ArgoCD, cert-manager"),
        ("app-pool", "m6i.2xlarge (8 vCPU, 32 GB)", "24", "All 47 business microservices"),
        ("data-pool", "r6i.4xlarge (16 vCPU, 128 GB)", "12", "Kafka, ClickHouse, PostgreSQL, Redis"),
        ("ml-pool", "g4dn.xlarge (4 vCPU, 16 GB, T4 GPU)", "6", "Triton inference, ML training"),
        ("spot-pool", "m6i.2xlarge (Spot)", "6", "Batch jobs, report generation, workers"),
    ]
    col_w3 = [25, 68, 15, 62]
    table_row([r[0] for r in node_rows[:1]], col_w3, bold=True)
    for row in node_rows[1:]:
        table_row(list(row), col_w3)
    pdf.ln(3)

    h2("5.3 Estimated Cloud Costs")
    body(
        "CONFLICT: The following cost estimates differ between the CTO and CFO versions of the budget. "
        "CTO estimate (for reference only - not approved by CFO): $31,700/month for cloud (dev+staging+prod). "
        "CFO approved budget: $480,000 total for 18 months = $26,667/month. "
        "CEO requested: $720,000 for 24 months = $30,000/month."
    )
    cost_rows = [
        ("Environment", "Monthly Cost (USD)", "Components"),
        ("Production (multi-AZ)", "$18,500", "48 nodes, load balancers, managed DBs, CDN"),
        ("Staging", "$6,200", "12 nodes, smaller DB instances"),
        ("Development", "$4,500", "8 nodes, shared DBs"),
        ("Data lake (S3 + Glacier)", "$1,200", "Raw data, backups, ML artifacts"),
        ("CDN (CloudFront)", "$800", "Static assets, API cache"),
        ("Misc (DNS, IAM, etc.)", "$500", "Route53, ACM, CloudWatch"),
        ("TOTAL", "$31,700 / month", "All environments combined"),
    ]
    col_w4 = [42, 40, 88]
    table_row([r[0] for r in cost_rows[:1]], col_w4, bold=True)
    for row in cost_rows[1:]:
        table_row(list(row), col_w4)
    pdf.ln(3)

    # Section 6: ML Platform
    h1("6. ML Platform")
    h2("6.1 Production Models")
    body(
        "The AI Engine (service: ai-engine) orchestrates 6 ML models in production. "
        "All models are versioned in MLflow, deployed via Triton Inference Server or FastAPI, "
        "and continuously monitored for data drift and accuracy degradation."
    )
    ml_rows = [
        ("Model", "Algorithm", "Latency Target", "Retrain Freq"),
        ("Demand Forecasting", "LSTM + Prophet", "< 500ms (batch)", "Weekly"),
        ("Churn Prediction", "XGBoost", "< 50ms (real-time)", "Monthly"),
        ("Anomaly Detection", "Isolation Forest", "< 20ms (stream)", "Weekly"),
        ("Route Optimization", "RL (PPO)", "< 2s (per request)", "Monthly"),
        ("Document Classification", "BERT-base", "< 200ms", "Quarterly"),
        ("Fraud Detection", "GNN", "< 100ms (real-time)", "Weekly"),
    ]
    col_w5 = [45, 32, 35, 28]
    table_row([r[0] for r in ml_rows[:1]], col_w5, bold=True)
    for row in ml_rows[1:]:
        table_row(list(row), col_w5)
    pdf.ln(3)

    h2("6.2 LLM Integration")
    body(
        "The BI & Analytics module uses GPT-4o (OpenAI API) for Natural Language Queries. "
        "Users type questions in Russian, Kazakh, or English; the system generates a ClickHouse SQL query, "
        "executes it, and returns structured results with a natural language explanation."
    )
    body(
        "Cost estimate per NLQ: $0.02-$0.05 depending on query complexity. "
        "Monthly cost estimate at 50,000 queries/day: $30,000-$75,000. "
        "This cost is included in the AI Pack add-on ($990-$1,200/month) amortised across tenants."
    )
    body(
        "Local fallback: Llama 3 8B via Ollama, used in air-gapped RU/KZ deployments. "
        "Quality trade-off: approximately 20% accuracy reduction on complex analytical queries."
    )

    # Section 7: Disaster Recovery
    h1("7. Disaster Recovery & Business Continuity")
    h2("7.1 Recovery Objectives by Tier")
    dr_rows = [
        ("Tier", "RTO", "RPO", "Strategy"),
        ("Enterprise", "15 min", "5 min", "Hot standby (multi-AZ) + async replication"),
        ("Business", "1 hour", "30 min", "Warm standby + PITR"),
        ("Starter", "4 hours", "1 hour", "Cold standby + daily backup restore"),
    ]
    col_w6 = [28, 20, 20, 102]
    table_row([r[0] for r in dr_rows[:1]], col_w6, bold=True)
    for row in dr_rows[1:]:
        table_row(list(row), col_w6)
    pdf.ln(3)

    h2("7.2 Backup Strategy")
    backup_items = [
        "PostgreSQL: continuous WAL streaming to S3 (PITR up to 7 days) + weekly full dump.",
        "ClickHouse: daily incremental + weekly full backup to S3.",
        "Kafka: MirrorMaker 2 replication to standby cluster in separate AZ.",
        "Redis: RDB snapshot every 1 hour; AOF for <= 1 min RPO on Enterprise.",
        "Elasticsearch: S3 snapshot every 6 hours.",
        "S3 object store: cross-region replication enabled (versioning ON).",
    ]
    for item in backup_items:
        bullet(item)

    # Section 8: Compliance
    h1("8. Regulatory Compliance")
    compliance_data = [
        ("Russia 152-FZ", "Data localisation: all RU personal data stored in Yandex Cloud ru-central1. "
         "Air-gapped from AWS. FSTEC certification planned Q2 2026."),
        ("Kazakhstan PDP 2023", "Personal data stored in Kazakhtelecom Cloud (Almaty). "
         "Planned Q3 2026 deployment. Currently deferred."),
        ("GDPR", "EU customer data in AWS eu-central-1 (Frankfurt). "
         "Data retention policies automated. DPA signed with AWS. "
         "Right to erasure: automated within 30 days of request."),
        ("PCI DSS Level 1", "Applicable to SwiftPay integration and billing-service. "
         "Annual QSA audit required. Tokenisation via Stripe (no raw card data stored)."),
        ("ISO 27001", "Certification planned Q3 2026. Gap assessment completed Q4 2024. "
         "12 non-conformities identified, remediation in progress."),
        ("ISO 9001:2015", "Quality management standard enforced in manufacturing module. "
         "Inspection records, NCR tracking, corrective actions fully digitised."),
    ]
    for standard, description in compliance_data:
        h3(standard)
        body(description)

    # Budget summary (CONFLICT)
    h1("9. Budget Summary (Technical Perspective)")
    body(
        "WARNING: The following figures are from the CTO's technical budget estimate (Section D of budget_draft.txt). "
        "These figures CONFLICT with both the CEO's version ($6,800,000 total) and the "
        "CFO's approved budget ($4,200,000 total). Final budget is pending board decision."
    )
    budget_rows = [
        ("Cost Category", "Monthly (USD)", "20-Month Total"),
        ("Development team salaries", "$164,000", "$3,280,000"),
        ("Cloud infrastructure (all envs)", "$31,700", "$634,000"),
        ("Tooling licenses (GitHub, Jira, etc.)", "$7,200", "$144,000"),
        ("Security & compliance", "$14,200", "$284,000"),
        ("Marketing & sales", "N/A", "$350,000"),
        ("Reserve (10%)", "Varies", "$469,200"),
        ("TOTAL (CTO estimate)", "$217,100", "$5,161,200"),
    ]
    col_w7 = [65, 45, 60]
    table_row([r[0] for r in budget_rows[:1]], col_w7, bold=True)
    for row in budget_rows[1:]:
        table_row(list(row), col_w7)
    pdf.ln(3)
    body(
        "Note: The CTO estimate of $5,161,200 falls between the CEO's requested $6,800,000 and "
        "the CFO's approved $4,200,000. All three figures are live and unresolved as of Feb 2025."
    )

    # Appendix
    h1("Appendix: Document History")
    hist_rows = [
        ("Version", "Date", "Author", "Changes"),
        ("1.0", "2024-11-01", "Denis Krasnov", "Initial draft"),
        ("1.5", "2024-12-10", "Denis Krasnov", "Added security section"),
        ("2.0", "2025-01-15", "Kirill Osipov", "Full rewrite with ML section"),
        ("2.1", "2025-02-20", "Denis Krasnov", "Added cost table, updated service count"),
    ]
    col_w8 = [20, 26, 40, 84]
    table_row([r[0] for r in hist_rows[:1]], col_w8, bold=True)
    for row in hist_rows[1:]:
        table_row(list(row), col_w8)

    out_path = os.path.join(OUT_DIR, "technical_specification.pdf")
    pdf.output(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    make_docx()
    make_pdf()
    print("Done.")
