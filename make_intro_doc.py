"""
Generates: NeuralDocs_AI_Intro.docx
Run: python make_intro_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "NeuralDocs_AI_Intro.docx"
ACCENT = RGBColor(0x1A, 0x73, 0xE8)   # blue
MUTED  = RGBColor(0x88, 0x88, 0x88)   # grey
PLACEHOLDER_BG = "FFF3CD"             # warm yellow fill for screenshot boxes

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name  = "Arial"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size  = 18 if level == 1 else 13
    color = ACCENT if level == 1 else RGBColor(0x33, 0x33, 0x33)
    set_font(run, size, bold=True, color=color)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        set_font(rb, 11, bold=True)
    r = p.add_run(text)
    set_font(r, 11)
    return p

def screenshot_box(label):
    """Yellow placeholder box with screenshot label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Yellow shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  PLACEHOLDER_BG)
    pPr.append(shd)

    # Border around the box
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "E6A817")
        pBdr.append(el)
    pPr.append(pBdr)

    run = p.add_run(f"  📸  СКРИНШОТ: {label}  ")
    set_font(run, 10, bold=True, color=RGBColor(0x85, 0x60, 0x00), italic=True)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "DDDDDD")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

# ═════════════════════════════════════════════════════════════════════════════
# TITLE
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("NeuralDocs AI")
set_font(r, 26, bold=True, color=ACCENT)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("Умный анализ проектной документации")
set_font(r2, 13, italic=True, color=MUTED)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(14)
r3 = p3.add_run("Краткое руководство пользователя")
set_font(r3, 10, color=MUTED)

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 1. ЧТО ЭТО
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Что такое NeuralDocs AI?")
body(
    "NeuralDocs AI — веб-приложение, которое автоматически анализирует проектные файлы "
    "и превращает хаотичные документы (черновики, конфиги, код, заметки) в структурированную "
    "спецификацию проекта. Система сама находит противоречия между документами — "
    "например, когда в одном файле бюджет $4 200 000, а в другом $6 800 000."
)

bullet("Поддерживаемые форматы:", bold_prefix="📄")
for fmt in ["TXT, MD (текстовые заметки и описания)",
            "JSON (конфигурационные файлы)",
            "PY, Java (исходный код)",
            "DOCX (документы Word)",
            "PDF (технические спецификации)"]:
    bullet(fmt)

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 2. КАК ЗАПУСТИТЬ
# ═════════════════════════════════════════════════════════════════════════════
heading("2. Запуск приложения")
body("Приложение запускается через Docker. Для этого нужно:")

bullet("Убедиться, что Docker Desktop запущен.", bold_prefix="1.")
bullet("Запустить LM Studio и включить локальный сервер (порт 1234).", bold_prefix="2.")
bullet("Выполнить команды в терминале:", bold_prefix="3.")

p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(1.5)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F4F4F4")
pPr.append(shd)
r = p.add_run("docker build -t omnicore-app .\ndocker run -p 8000:8000 --add-host=host.docker.internal:host-gateway omnicore-app")
r.font.name = "Courier New"
r.font.size = Pt(10)

bullet("Открыть браузер и перейти по адресу http://localhost:8000", bold_prefix="4.")

screenshot_box("Главная страница приложения в браузере")

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 3. КАК ПОЛЬЗОВАТЬСЯ
# ═════════════════════════════════════════════════════════════════════════════
heading("3. Как анализировать документы")

body("Весь процесс занимает несколько кликов:")

steps = [
    ("Загрузите файлы",
     "Нажмите на область загрузки или перетащите файлы проекта. "
     "Можно загружать сразу несколько файлов разных форматов."),
    ("Нажмите «Анализировать»",
     "Система обработает каждый файл через локальную языковую модель "
     "и извлечёт ключевые факты."),
    ("Дождитесь результата",
     "Время обработки зависит от объёма файлов и мощности компьютера. "
     "Обычно занимает от 1 до 10 минут."),
]
for title, desc in steps:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    rb = p.add_run(title + " — ")
    set_font(rb, 11, bold=True)
    r = p.add_run(desc)
    set_font(r, 11)

screenshot_box("Интерфейс загрузки файлов и кнопка «Анализировать»")

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 4. ЧТО ПОЛУЧАЕТСЯ
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Результат анализа")
body(
    "После обработки система выдаёт структурированный документ со следующими полями:"
)

fields = [
    ("Обзор проекта",         "краткое описание из документов"),
    ("Цели",                  "список целей с указанием источника"),
    ("Требования",            "функциональные и технические требования"),
    ("Технологический стек",  "языки, фреймворки, базы данных"),
    ("Архитектура",           "архитектурные решения"),
    ("Команда",               "участники и их роли"),
    ("Сроки",                 "дедлайны и этапы"),
    ("Бюджет",                "суммы с указанием источника"),
    ("Риски",                 "выявленные риски проекта"),
]
for field, desc in fields:
    bullet(f"— {desc}", bold_prefix=field)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
rb = p.add_run("Обнаружение противоречий: ")
set_font(rb, 11, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
r = p.add_run(
    "если в разных файлах одно и то же поле (например, бюджет или срок) "
    "содержит разные значения — система выделит конфликт красным и покажет, "
    "в каком файле и строке найдено каждое значение."
)
set_font(r, 11)

screenshot_box("Результат анализа — структурированная спецификация с выделенными конфликтами")
screenshot_box("Детали конфликта — сравнение двух противоречащих значений")

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 5. АРХИВ
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Архив готовых примеров")
body(
    "Во вкладке «Архив» можно просмотреть 100 готовых примеров анализа для 10 вымышленных "
    "компаний. Это удобно для знакомства с системой без необходимости загружать собственные файлы."
)
bullet("Поиск по названию проекта или ключевым словам", bold_prefix="🔍")
bullet("Фильтрация по количеству найденных конфликтов", bold_prefix="⚠️")
bullet("Просмотр полной спецификации одним кликом", bold_prefix="📋")
bullet("Скачивание результата в формате JSON", bold_prefix="💾")

screenshot_box("Вкладка «Архив» со списком готовых примеров")

divider()

# ═════════════════════════════════════════════════════════════════════════════
# 6. ВАЖНЫЕ ЗАМЕЧАНИЯ
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Важные замечания")
bullet(
    "Все вычисления выполняются локально на вашем компьютере. "
    "Никакие данные не передаются в интернет.",
    bold_prefix="🔒 Конфиденциальность:"
)
bullet(
    "Качество анализа зависит от модели. Рекомендуется использовать "
    "meta-llama-3-8b-instruct в LM Studio.",
    bold_prefix="🤖 Модель:"
)
bullet(
    "При большом объёме файлов обработка занимает больше времени. "
    "Рекомендуется загружать не более 5–7 файлов за раз.",
    bold_prefix="⏱ Производительность:"
)
bullet(
    "Язык интерфейса: русский, казахский, английский.",
    bold_prefix="🌐 Язык:"
)

# ═════════════════════════════════════════════════════════════════════════════
# FOOTER note
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NeuralDocs AI — анализ документов на базе локальных LLM-моделей")
set_font(r, 9, italic=True, color=MUTED)

doc.save(OUT)
print(f"Готово: {OUT}")
