import os
import io
import re
import ast
import time
import json
from typing import List
from concurrent.futures import ThreadPoolExecutor
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import HTMLResponse, JSONResponse
from openai import OpenAI


# ─────────────────────────────────────────────────────────────────────────────
# FILE PARSERS (V2 Change 1): extract real signal from binary / code files
# before feeding them to the LLM. Cuts token cost dramatically on large
# projects (e.g. 11_omnicore_platform: ~190 KB → ~50-60 KB).
# ─────────────────────────────────────────────────────────────────────────────
def _parse_docx(content: bytes) -> str:
    """Extract paragraphs + table cell text from a .docx file."""
    try:
        from docx import Document
    except ImportError:
        return content.decode("utf-8", errors="ignore")
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        return f"[DOCX parse error: {e}]"
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _parse_pdf(content: bytes) -> str:
    """Extract text per page from a PDF."""
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # fallback
        except ImportError:
            return "[PDF parser not installed]"
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as e:
        return f"[PDF parse error: {e}]"
    out = []
    for i, page in enumerate(reader.pages, 1):
        try:
            t = (page.extract_text() or "").strip()
        except Exception:
            t = ""
        if t:
            out.append(f"--- page {i} ---\n{t}")
    return "\n\n".join(out)


def _parse_python(source: str) -> str:
    """
    AST-based extraction: module docstring, top-level constants,
    class/function signatures with their docstrings. Skips function bodies.
    Falls back to raw comments if AST fails.
    """
    out = []
    try:
        tree = ast.parse(source)
    except Exception:
        # fallback: keep only comments and docstrings-ish content
        keep = [ln for ln in source.splitlines()
                if ln.strip().startswith("#") or '"""' in ln or "'''" in ln]
        return "\n".join(keep)

    mod_doc = ast.get_docstring(tree)
    if mod_doc:
        out.append(f'"""{mod_doc}"""')

    for node in tree.body:
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            try:
                out.append(ast.unparse(node))
            except Exception:
                pass
        elif isinstance(node, ast.Assign):
            # top-level constants (UPPER_CASE)
            for tgt in node.targets:
                if isinstance(tgt, ast.Name) and tgt.id.isupper():
                    try:
                        out.append(ast.unparse(node))
                    except Exception:
                        pass
                    break
        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            try:
                args = ast.unparse(node.args)
            except Exception:
                args = "..."
            out.append(f"def {node.name}({args}):")
            d = ast.get_docstring(node)
            if d:
                out.append(f'    """{d}"""')
        elif isinstance(node, ast.ClassDef):
            bases = ""
            try:
                bases = ", ".join(ast.unparse(b) for b in node.bases)
            except Exception:
                pass
            out.append(f"class {node.name}({bases}):")
            d = ast.get_docstring(node)
            if d:
                out.append(f'    """{d}"""')
            for sub in node.body:
                if isinstance(sub, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    try:
                        args = ast.unparse(sub.args)
                    except Exception:
                        args = "..."
                    out.append(f"    def {sub.name}({args}):")
                    sd = ast.get_docstring(sub)
                    if sd:
                        out.append(f'        """{sd}"""')
    return "\n".join(out)


_JAVA_SIG_RE = re.compile(
    r"^\s*(?:public|private|protected|static|final|abstract|synchronized|\s)*"
    r"[\w<>\[\],\s]+\s+\w+\s*\([^)]*\)\s*(?:throws[\w\s,]+)?\s*[{;]",
    re.MULTILINE,
)
_JAVA_CLASS_RE = re.compile(
    r"^\s*(?:public|private|protected|abstract|final|\s)*"
    r"(?:class|interface|enum)\s+\w+[^{]*\{",
    re.MULTILINE,
)


def _parse_java(source: str) -> str:
    """
    Regex-based extraction: package, imports, javadoc/comments,
    class/interface headers, method signatures. Skips bodies.
    """
    out = []
    # package + imports
    for m in re.finditer(r"^\s*(package|import)\s+[^;]+;", source, re.MULTILINE):
        out.append(m.group(0).strip())

    # javadoc /** ... */ blocks
    for m in re.finditer(r"/\*\*.*?\*/", source, re.DOTALL):
        out.append(m.group(0))

    # // comments that look informative (skip trivial 1-word ones)
    for m in re.finditer(r"//[^\n]{6,}", source):
        out.append(m.group(0).strip())

    # class / interface / enum headers (without body)
    for m in _JAVA_CLASS_RE.finditer(source):
        out.append(m.group(0).rstrip("{").strip())

    # method signatures (drop the opening brace/semicolon)
    for m in _JAVA_SIG_RE.finditer(source):
        sig = m.group(0).rstrip("{;").strip()
        # filter noise: skip if it's actually an `if (...)` or similar control
        if re.match(r"^\s*(if|for|while|switch|catch|return)\b", sig):
            continue
        out.append(sig)

    return "\n".join(out)


def _flatten_json(content: bytes) -> str:
    """
    Convert a JSON file into flat "key: value" lines so the MAP model
    receives simple strings instead of nested dicts/arrays.
    Example: {"tech_stack": ["Python", "FastAPI"]} → "tech_stack: Python\ntech_stack: FastAPI"
    """
    try:
        data = json.loads(content.decode("utf-8", errors="ignore"))
    except Exception:
        return content.decode("utf-8", errors="ignore")

    lines: list = []

    def _walk(obj, prefix: str = ""):
        if isinstance(obj, dict):
            for k, v in obj.items():
                new_prefix = f"{k}: " if not prefix else f"{prefix}{k}: "
                _walk(v, new_prefix)
        elif isinstance(obj, list):
            for item in obj:
                _walk(item, prefix)
        else:
            val = str(obj).strip()
            if val:
                lines.append(f"{prefix}{val}" if prefix else val)

    _walk(data)
    return "\n".join(lines)


def _parse_markdown(text: str) -> str:
    """
    Strip heavy non-prose content from Markdown files so the MAP model
    receives readable text instead of SQL, protobuf, ASCII art, and YAML.

    Keeps: section headers (##), bullet points, regular paragraphs.
    Removes:
      • Fenced code blocks  (``` ... ```)
      • Lines that are mostly box-drawing characters (ASCII art diagrams)
      • Markdown table divider rows (|---|---|)
      • Trailing whitespace / excessive blank lines
    """
    # 1. Remove fenced code blocks (``` ... ``` or ~~~ ... ~~~)
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"~~~[\s\S]*?~~~", "", text)

    lines_out = []
    prev_blank = False
    for line in text.splitlines():
        stripped = line.strip()

        # Skip table divider rows: |---|---|
        if re.match(r"^\|[\s\-\|:]+\|$", stripped):
            continue

        # Skip lines that are ≥ 40% box-drawing / ASCII-art characters
        art_chars = sum(1 for c in stripped if c in "│├└─┐┘┌┤┬┴┼╔╗╚╝║═╠╣╦╩╬▼▲►◄")
        if stripped and art_chars / len(stripped) >= 0.4:
            continue

        # Collapse multiple blank lines into one
        if stripped == "":
            if not prev_blank:
                lines_out.append("")
            prev_blank = True
        else:
            lines_out.append(line.rstrip())
            prev_blank = False

    return "\n".join(lines_out).strip()


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch a file to the right parser based on extension."""
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            return _parse_docx(content)
        if name.endswith(".pdf"):
            return _parse_pdf(content)
        # textual formats from here on
        text = content.decode("utf-8", errors="ignore")
        if name.endswith(".py"):
            # Prefix tells the MAP model this is source code, not a spec document
            return "[SOURCE CODE — extract only technical_solution/architecture]\n" + _parse_python(text)
        if name.endswith(".java"):
            return "[SOURCE CODE — extract only technical_solution/architecture]\n" + _parse_java(text)
        if name.endswith(".json"):
            # Flatten nested JSON to "key: value" lines — prevents MAP from
            # copying complex objects verbatim as fact values
            return _flatten_json(content)
        if name.endswith(".md"):
            # Strip code blocks and ASCII art — large .md files like architecture.md
            # contain massive SQL/protobuf/YAML sections that overwhelm the model
            return _parse_markdown(text)
        # .txt and anything else → raw text
        return text
    except Exception as e:
        return f"[parser error for {filename}: {e}]"

# Path to the pre-generated samples archive (in the same folder as this file)
SAMPLES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "finetune_v2")

class LLMProvider:
    def generate_document(self, files_data: list[dict]) -> dict:
        pass

# --- ЛОКАЛЬНЫЙ ПРОВАЙДЕР С ЖЕСТКОЙ ДЕТЕРМИНИРОВАННОСТЬЮ ---
class LocalProvider(LLMProvider):
    def __init__(self, base_url: str, model_name: str = "meta-llama-3-8b-instruct"):
        self.client = OpenAI(base_url=base_url, api_key="lm-studio-local")
        self.model_name = model_name

    def generate_document(self, files_data: list[dict], language: str = "ru") -> dict:
        start_time = time.time()
        total_tokens_used = 0

        # ── Change 2: Schema-guided MAP ───────────────────────────────────────
        # MAP now outputs a compact JSON keyed by the 9 schema fields instead of
        # free-form "[N] Факт: ..." lines.  Python merges all chunk JSONs, detects
        # budget/timeline conflicts deterministically, then feeds a concise
        # structured fact list to the REDUCE LLM call.  Benefits:
        #  • Model only writes facts that map to a known field → zero noise tokens
        #  • Conflict detection is exact (Python string comparison), not guessed
        #  • REDUCE prompt is shorter → more room for the actual content
        MAP_CHUNK_CHARS  = 4_000   # chars per MAP call (≈ 1 500 tokens) — smaller
        #                            gives json_schema grammar more VRAM headroom
        REDUCE_MAX_CHARS = 5_000   # merged fact list cap for REDUCE

        SCHEMA_FIELDS_LIST   = ["goals", "requirements", "team", "risks"]
        SCHEMA_FIELDS_SCALAR = ["technical_solution", "architecture", "timeline", "budget"]

        map_instruction = """\
You are a structured fact extractor. Read the document chunk and output ONLY a JSON object \
with the fields below. Include ONLY fields you find clear evidence for — omit the rest entirely.

{
  "goals":              ["<project goal or objective — copy text exactly>"],
  "requirements":       ["<functional or technical requirement — copy text exactly>"],
  "technical_solution": "<programming languages, frameworks, engines, databases — copy exactly>",
  "architecture":       "<system design, components, deployment — copy exactly>",
  "team":               ["<any person, role, or team member mentioned — copy exactly>"],
  "timeline":           "<any duration, deadline, or timeframe mentioned — copy exactly>",
  "budget":             "<any monetary amount or cost mentioned — copy exactly>",
  "risks":              ["<any risk, problem, or concern mentioned — copy exactly>"]
}

Rules:
• Copy text EXACTLY as it appears, in the original language. Do not translate.
• If a field has no evidence in this chunk, omit that key completely.
• Output ONLY the JSON object — no markdown fences, no commentary.
• IMPORTANT: If the chunk starts with [SOURCE CODE], only extract technical_solution \
and architecture from it. Function names and method signatures are NOT goals or requirements."""

        def _parse_map_output(raw: str) -> dict:
            """Try to parse the MAP JSON; return {} on any failure."""
            s = raw.find("{"); e = raw.rfind("}")
            if s == -1 or e == -1:
                return {}
            snippet = raw[s:e+1]
            try:
                return json.loads(snippet)
            except Exception:
                # Strip trailing commas (common small-model mistake) and retry
                cleaned = re.sub(r",\s*([}\]])", r"\1", snippet)
                try:
                    return json.loads(cleaned)
                except Exception:
                    return {}

        # ── Structured-output schemas ──────────────────────────────────────────
        # LM Studio supports OpenAI-compatible response_format.
        # "json_object"  → llama.cpp JSON grammar: always valid JSON, no fences.
        # "json_schema"  → token-level GBNF grammar: forces exact field names +
        #                  types (LM Studio 0.3.5+).  We try json_schema first for
        #                  REDUCE and fall back to json_object if the server rejects
        #                  it (older LM Studio versions).

        # MAP schema: all fields optional (model omits keys it has no evidence for)
        _MAP_RESPONSE_FORMAT = {"type": "json_object"}

        # Shared shape for a single fact entry (used in list + scalar fields)
        _FACT_ITEM = {
            "type": "object",
            "properties": {
                "text":             {"type": "string"},
                "source":           {"type": "string"},
                "has_conflict":     {"type": "boolean"},
                "conflict_details": {"type": "string"}
            },
            "required": ["text", "source", "has_conflict", "conflict_details"],
            "additionalProperties": False
        }

        # REDUCE schema: all 9 fields always present, correct types enforced
        _REDUCE_RESPONSE_FORMAT = {
            "type": "json_schema",
            "json_schema": {
                "name": "project_spec",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "project_overview":   {"type": "string"},
                        "goals":              {"type": "array",  "items": _FACT_ITEM},
                        "requirements":       {"type": "array",  "items": _FACT_ITEM},
                        "technical_solution": _FACT_ITEM,
                        "architecture":       _FACT_ITEM,
                        "team":               {"type": "array",  "items": _FACT_ITEM},
                        "timeline":           _FACT_ITEM,
                        "budget":             _FACT_ITEM,
                        "risks":              {"type": "array",  "items": _FACT_ITEM}
                    },
                    "required": [
                        "project_overview", "goals", "requirements",
                        "technical_solution", "architecture",
                        "team", "timeline", "budget", "risks"
                    ],
                    "additionalProperties": False
                }
            }
        }

        def _map_chunk(filename, lines, line_offset):
            """Send one chunk to the LLM; return (parsed_dict, token_count)."""
            numbered = "\n".join(
                f"[{line_offset+i+1}] {ln}" for i, ln in enumerate(lines)
            )
            try:
                resp = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": map_instruction},
                        {"role": "user",   "content": f"File: {filename}\n\n{numbered}"}
                    ],
                    temperature=0.0,
                    response_format=_MAP_RESPONSE_FORMAT
                )
            except Exception:
                # Fallback: older LM Studio / model doesn't support response_format
                resp = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": map_instruction},
                        {"role": "user",   "content": f"File: {filename}\n\n{numbered}"}
                    ],
                    temperature=0.0
                )
            tokens = resp.usage.total_tokens if resp.usage else 0
            parsed = _parse_map_output(resp.choices[0].message.content)
            # Normalise: wrap accidental scalar in list where list is expected
            for key in SCHEMA_FIELDS_LIST:
                if key in parsed and isinstance(parsed[key], str):
                    parsed[key] = [parsed[key]]
            return parsed, tokens

        def _map_one(f):
            lines = f["content"].splitlines()
            # Split into context-safe chunks
            chunks, cur, cur_len, offset = [], [], 0, 0
            for i, line in enumerate(lines):
                ll = len(line) + 1
                if cur_len + ll > MAP_CHUNK_CHARS and cur:
                    chunks.append((cur, offset))
                    offset = i
                    cur, cur_len = [], 0
                cur.append(line)
                cur_len += ll
            if cur:
                chunks.append((cur, offset))

            file_dicts, total_tokens = [], 0
            for idx, (chunk_lines, line_offset) in enumerate(chunks):
                label = (f"{f['filename']} (part {idx+1}/{len(chunks)})"
                         if len(chunks) > 1 else f['filename'])
                d, tok = _map_chunk(label, chunk_lines, line_offset)
                if d:
                    d["_source"] = label
                file_dicts.append(d)
                total_tokens += tok
            return f["filename"], file_dicts, total_tokens

        # ── Run MAP sequentially (avoids OOM on CPU inference) ────────────────
        max_workers = 1
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            map_results = list(pool.map(_map_one, files_data))

        # ── Python merge: combine all chunk dicts ─────────────────────────────
        # List fields → extend; scalar fields → collect all distinct values
        # (multiple distinct budget/timeline values = conflict evidence)
        merged: dict = {k: [] for k in SCHEMA_FIELDS_LIST + SCHEMA_FIELDS_SCALAR}
        for filename, chunk_dicts, tokens in map_results:
            total_tokens_used += tokens
            for d in chunk_dicts:
                src = d.get("_source", filename)
                for key in SCHEMA_FIELDS_LIST:
                    if key in d and isinstance(d[key], list):
                        for item in d[key]:
                            if item and str(item).strip():
                                merged[key].append({"text": str(item), "source": src})
                for key in SCHEMA_FIELDS_SCALAR:
                    if key in d and d[key]:
                        val = d[key]
                        # MAP sometimes returns a list for a scalar field
                        # (e.g. when the model finds multiple values in the chunk).
                        # Join them into a single string rather than repr-ing the list.
                        if isinstance(val, list):
                            val = "; ".join(
                                str(v).strip() for v in val if str(v).strip()
                            )
                        elif not isinstance(val, str):
                            val = str(val)
                        val = val.strip()
                        if val:
                            merged[key].append({"text": val, "source": src})

        # ── Change 4: Regex pre-extraction ────────────────────────────────────
        # Scan every file's parsed text with regex patterns for the three fields
        # most likely to be missed or garbled by the LLM (budget, timeline, team).
        # Regex hits are added to `merged` so Python conflict detection sees them
        # alongside whatever the MAP phase found.  No LLM call — 100% reliable
        # for clearly formatted numeric/structured values.

        # Normalise a monetary string: collapse whitespace, unify currency symbols
        def _norm_money(s: str) -> str:
            s = re.sub(r"\s+", " ", s.strip())
            s = re.sub(r"тнг\b", "тенге", s, flags=re.IGNORECASE)
            return s

        # Regex patterns — each tuple: (compiled_pattern, group_index, field)
        BUDGET_RES = [
            # "30 000 000 тенге" / "30000000 ₸"  — require at least 4 digits total
            re.compile(r"\b(\d[\d\s]{3,}\s*(?:тенге|тнг|₸))", re.IGNORECASE),
            # "$4,200,000" / "$ 4 200 000" — stop before ". N." list-item suffixes
            re.compile(r"(\$\s*\d[\d\s]*\d)(?!\s*[,\s]*\d{3,})(?=\D|$)", re.IGNORECASE),
            # "4 200 000 usd/eur/руб"
            re.compile(r"\b(\d[\d\s]{3,}\s*(?:usd|eur|руб\.?|rub))\b", re.IGNORECASE),
            # "бюджет: X" / "budget: X" — grab rest of line only if it contains a digit
            re.compile(r"(?:бюджет|budget)\s*[:\-]\s*(\d[^\n]{2,40})", re.IGNORECASE),
        ]

        TIMELINE_RES = [
            # "20 месяцев" / "6 мес." / "12 months"
            re.compile(r"\b(\d+\s*(?:месяц(?:ев|а)?|мес\.?|month|months|week|weeks))\b", re.IGNORECASE),
            # "Q2 2026" / "Q4 2025"
            re.compile(r"\b(Q[1-4]\s*20\d{2})\b", re.IGNORECASE),
            # ISO date "2026-03-15" or "15.03.2026"
            re.compile(r"\b(20\d{2}-\d{2}-\d{2})\b"),
            re.compile(r"\b(\d{2}\.\d{2}\.20\d{2})\b"),
            # "через X месяцев / лет"
            re.compile(r"(через\s+\d+\s*(?:месяц\w*|нед\w*|год\w*|лет))", re.IGNORECASE),
        ]

        TEAM_RES = [
            # "2 iOS-разработчика" / "3 developer" / "5 engineers"
            re.compile(
                r"\b(\d+\s*(?:iOS|Android|бэкенд|backend|frontend|разработчик\w*|"
                r"developer|engineer|designer|qa|тестировщик\w*|продакт\w*|аналитик\w*|"
                r"devops|архитектор\w*))",
                re.IGNORECASE,
            ),
            # "Имя: Роль" or "Name — role" informal lines
            re.compile(r"^[-•*]\s*(.{5,60})\s*[:\-–—]\s*(.{3,40})$", re.MULTILINE),
        ]

        already_in_merged: dict = {
            "budget":   {re.sub(r"\s+", " ", e["text"].lower()) for e in merged["budget"]},
            "timeline": {re.sub(r"\s+", " ", e["text"].lower()) for e in merged["timeline"]},
            "team":     {re.sub(r"\s+", " ", e["text"].lower()) for e in merged["team"]},
        }

        for f in files_data:
            fname = f["filename"]
            text  = f["content"]

            # Budget
            for pat in BUDGET_RES:
                for m in pat.finditer(text):
                    val = _norm_money(m.group(1))
                    norm = re.sub(r"\s+", " ", val.lower())
                    # Skip if already present or too short / looks like noise
                    if norm in already_in_merged["budget"] or len(val) < 4:
                        continue
                    # Skip if it's clearly not a monetary value (e.g. plain word from бюджет: line)
                    if not re.search(r"\d", val):
                        continue
                    merged["budget"].append({"text": val, "source": fname})
                    already_in_merged["budget"].add(norm)

            # Timeline
            for pat in TIMELINE_RES:
                for m in pat.finditer(text):
                    val = m.group(1).strip()
                    norm = re.sub(r"\s+", " ", val.lower())
                    if norm in already_in_merged["timeline"] or len(val) < 3:
                        continue
                    merged["timeline"].append({"text": val, "source": fname})
                    already_in_merged["timeline"].add(norm)

            # Team
            for pat in TEAM_RES:
                for m in pat.finditer(text):
                    # Two-group pattern (name — role): groups == 2
                    if pat.groups == 2 and m.lastindex and m.lastindex >= 2:
                        try:
                            val = f"{m.group(1).strip()}: {m.group(2).strip()}"
                        except Exception:
                            val = m.group(0).strip()
                    else:
                        val = m.group(1).strip()
                    norm = re.sub(r"\s+", " ", val.lower())
                    if norm in already_in_merged["team"] or len(val) < 4:
                        continue
                    merged["team"].append({"text": val, "source": fname})
                    already_in_merged["team"].add(norm)

        # ── Python conflict detection (deterministic) ─────────────────────────
        # A real conflict = DIFFERENT VALUES in DIFFERENT FILES.
        # Multiple values from the same file are budget line items or milestone
        # dates — NOT contradictions.  We strip " (part N/M)" suffixes so that
        # different chunks of the same file are treated as the same source.
        def _base_src(source: str) -> str:
            """Return bare filename, stripping chunk labels like '(part 2/4)'."""
            return re.sub(r"\s*\(part\s+\d+/\d+\)\s*$", "", source.strip())

        def _detect_conflict(entries: list) -> tuple:
            if len(entries) < 2:
                return False, ""

            # Keep one representative (longest) value per base file
            per_file: dict = {}   # base_src → best entry
            for e in entries:
                src = _base_src(e["source"])
                existing = per_file.get(src)
                if existing is None or len(e["text"]) > len(existing["text"]):
                    per_file[src] = e

            if len(per_file) < 2:
                return False, ""   # all values from the same file → no conflict

            # Check whether the per-file values actually differ
            norms = {re.sub(r"\s+", " ", e["text"].strip().lower())
                     for e in per_file.values()}
            if len(norms) < 2:
                return False, ""   # all files agree → no conflict

            parts = "; ".join(
                f'{src} — "{e["text"]}"' for src, e in per_file.items()
            )
            return True, f"Conflict: {parts}"

        budget_conflict,   budget_detail   = _detect_conflict(merged["budget"])
        timeline_conflict, timeline_detail = _detect_conflict(merged["timeline"])

        # ── Deduplicate scalar fields: one best value per source file ──────────
        # The regex extraction can add 5-6 budget line-items all from the same
        # file.  Passing all of them to REDUCE overwhelms the model and causes it
        # to write a complex dict/list as the budget value → broken JSON.
        # We keep the LONGEST (most informative) entry per base filename.
        for key in SCHEMA_FIELDS_SCALAR:
            if len(merged[key]) <= 1:
                continue
            per_file: dict = {}
            for e in merged[key]:
                src = _base_src(e["source"])
                cur = per_file.get(src)
                if cur is None or len(e["text"]) > len(cur["text"]):
                    per_file[src] = e
            merged[key] = list(per_file.values())

        # ── Build compact fact summary ─────────────────────────────────────────
        # Format: FIELD | source | value  (one line per entry)
        fact_lines = []
        for key in SCHEMA_FIELDS_LIST + SCHEMA_FIELDS_SCALAR:
            for e in merged[key]:
                fact_lines.append(f"{key} | {e['source']} | {e['text']}")
        all_extracted_facts = "\n".join(fact_lines)   # kept for hallucination check corpus

        # ── Change 3: Hierarchical REDUCE ─────────────────────────────────────
        # If the merged fact list fits in one REDUCE call → single pass (fast).
        # If it overflows (large projects like 11_omnicore_platform) → split into
        # groups, run one REDUCE LLM call per group to produce partial specs, then
        # merge the partial specs in Python.  No extra LLM call for the merge —
        # conflicts are already detected deterministically by Python above.

        conflict_hints = ""
        if budget_conflict:
            conflict_hints += f"\nBUDGET CONFLICT: {budget_detail}"
        if timeline_conflict:
            conflict_hints += f"\nTIMELINE CONFLICT: {timeline_detail}"

        json_template = """{
  "project_overview": "<1-2 sentence summary in Russian of what the project is>",
  "goals":        [{"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""}],
  "requirements": [{"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""}],
  "technical_solution": {"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""},
  "architecture":       {"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""},
  "team":    [{"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""}],
  "timeline":     {"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""},
  "budget":       {"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""},
  "risks":   [{"text": "<fact text>", "source": "<source>", "has_conflict": false, "conflict_details": ""}]
}"""

        def _build_reduce_instruction(facts_text: str, include_conflicts: bool = True) -> str:
            hints = conflict_hints if include_conflicts else ""
            return f"""You are a strict JSON generator. Fill the template below using ONLY the FACTS provided.

FACTS format — each line: FIELD | source | value

RULES:
• Use ONLY the facts listed. Never invent, translate, or paraphrase beyond the given text.
• Keep all text in the SAME LANGUAGE as the source fact (do not translate Russian to English).
• If a field has no facts, write "Нет данных" in "text" and "" in "source".
• Never leave "text" as an empty string "".
• For list fields (goals, requirements, team, risks): one item per unique fact.
• For scalar fields (technical_solution, architecture, timeline, budget):
  - "text" must contain the ACTUAL VALUE (e.g. "30 000 000 тенге" or "6 мес"), NOT a conflict description.
  - copy the source of that value into "source".
• project_overview: write 1-2 sentences in Russian summarising what the project is about.
• Output ONLY the JSON object — no markdown, no commentary.

CONFLICT HINTS (pre-detected by the system):{hints if hints else " none"}
• If a conflict hint exists for budget/timeline → set has_conflict=true and paste the HINT TEXT into conflict_details. The "text" field must still contain the actual value.
• All other fields → has_conflict=false, conflict_details="".

TEMPLATE (replace every <fact text>/<source> with real values from FACTS):
{json_template}
"""

        def _call_reduce(facts_text: str, include_conflicts: bool = True) -> tuple:
            """One REDUCE LLM call. Returns (raw_response_str, token_count).
            Tries json_schema → json_object → plain text (in order of strictness).
            """
            instruction = _build_reduce_instruction(facts_text, include_conflicts)
            messages = [
                {"role": "system", "content": instruction},
                {"role": "user",   "content": f"FACTS:\n{facts_text}"}
            ]
            # Tier 1: json_schema — token-level grammar, exact field types enforced
            try:
                resp = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=messages,
                    temperature=0.0,
                    response_format=_REDUCE_RESPONSE_FORMAT
                )
                tokens = resp.usage.total_tokens if resp.usage else 0
                return resp.choices[0].message.content, tokens
            except Exception:
                pass
            # Tier 2: json_object — forces valid JSON but no schema constraint
            try:
                resp = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=messages,
                    temperature=0.0,
                    response_format={"type": "json_object"}
                )
                tokens = resp.usage.total_tokens if resp.usage else 0
                return resp.choices[0].message.content, tokens
            except Exception:
                pass
            # Tier 3: no format constraint — original behaviour, repair handles it
            resp = self.client.chat.completions.create(
                model=self.model_name,
                messages=messages,
                temperature=0.0
            )
            tokens = resp.usage.total_tokens if resp.usage else 0
            return resp.choices[0].message.content, tokens

        def _merge_partial_specs(dicts: list) -> dict:
            """
            Merge a list of partial parsed-spec dicts (output of intermediate
            REDUCE calls) into one combined dict — pure Python, no LLM call.
            List fields are concatenated; scalar fields take the first non-empty
            value (conflicts are stamped by the Python detection step later).
            """
            NON_DATA = {"нет данных", "no data", "данные отсутствуют", ""}

            result: dict = {
                "project_overview": "",
                **{k: [] for k in SCHEMA_FIELDS_LIST},
                **{k: {"text": "Нет данных", "source": "",
                       "has_conflict": False, "conflict_details": ""}
                   for k in SCHEMA_FIELDS_SCALAR},
            }

            for d in dicts:
                if not isinstance(d, dict) or "error" in d:
                    continue

                # project_overview — keep first useful one
                if not result["project_overview"]:
                    ov = (d.get("project_overview") or "").strip()
                    if ov and ov.lower() not in NON_DATA and "<" not in ov:
                        result["project_overview"] = ov

                # list fields — extend, skip empty placeholders
                for key in SCHEMA_FIELDS_LIST:
                    for item in (d.get(key) or []):
                        if not isinstance(item, dict):
                            continue
                        if (item.get("text") or "").strip().lower() in NON_DATA:
                            continue
                        result[key].append(item)

                # scalar fields — take first non-empty value
                for key in SCHEMA_FIELDS_SCALAR:
                    cur = result[key]
                    if cur.get("text", "").strip().lower() not in NON_DATA:
                        continue   # already filled
                    val = d.get(key)
                    if isinstance(val, dict):
                        vt = (val.get("text") or "").strip()
                        if vt and vt.lower() not in NON_DATA:
                            result[key] = val

            return result

        # ── Decide: single REDUCE or hierarchical ─────────────────────────────
        if len(all_extracted_facts) <= REDUCE_MAX_CHARS:
            # ── Fast path: fits in one call ───────────────────────────────────
            raw_content, tokens = _call_reduce(all_extracted_facts)
            total_tokens_used += tokens
        else:
            # ── Hierarchical path: split → partial REDUCEs → Python merge ─────
            lines = all_extracted_facts.split("\n")
            groups: list[list[str]] = []
            cur_grp: list[str] = []
            cur_len = 0
            for line in lines:
                ll = len(line) + 1
                if cur_len + ll > REDUCE_MAX_CHARS and cur_grp:
                    groups.append(cur_grp)
                    cur_grp, cur_len = [], 0
                cur_grp.append(line)
                cur_len += ll
            if cur_grp:
                groups.append(cur_grp)

            partial_dicts: list[dict] = []
            for i, grp_lines in enumerate(groups):
                grp_text = "\n".join(grp_lines)
                # Only pass conflict hints to the last group to avoid duplication
                raw_part, tokens = _call_reduce(
                    grp_text, include_conflicts=(i == len(groups) - 1)
                )
                total_tokens_used += tokens
                # Reuse _repair_json defined below — forward-declared via closure;
                # we'll parse after _repair_json is defined.
                partial_dicts.append(raw_part)   # store raw strings for now

            # _repair_json is defined in the helpers block below; parse after
            raw_content = None          # signal: use partial_dicts path
            _partial_raw_list = partial_dicts  # stash for post-helper use

        end_time = time.time()
        duration_ms = int((end_time - start_time) * 1000)

        # ── helpers (defined before use) ─────────────────────────────────────
        def _sanitize_placeholders(obj):
            """
            Safety net: if the LLM copied template placeholder words into
            conflict_details (e.g. 'срок_A', 'сумма_B', 'файл_A') instead of
            real extracted values, clear the field and reset has_conflict=False.
            Also clears has_conflict=True when conflict_details ended up empty.
            """
            import re
            # Matches placeholder patterns the LLM sometimes copies from the template
            PLACEHOLDER = re.compile(
                r'\b(?:срок|сумма|файл|значение|value|term|file|part)_[A-Za-zА-Яа-я0-9]{1,2}\b'
                r'|\.{3}',   # also catch literal "..." left in the template fields
                re.IGNORECASE
            )
            if not isinstance(obj, dict):
                return
            if 'has_conflict' in obj and 'conflict_details' in obj:
                details = obj.get('conflict_details') or ''
                if obj['has_conflict']:
                    if not details.strip() or PLACEHOLDER.search(details):
                        obj['has_conflict'] = False
                        obj['conflict_details'] = ''
            for val in obj.values():
                if isinstance(val, dict):
                    _sanitize_placeholders(val)
                elif isinstance(val, list):
                    for item in val:
                        if isinstance(item, dict):
                            _sanitize_placeholders(item)

        def _sanitize_hallucinations(obj, source_corpus: str):
            """
            Second pass: walk every {text, source} leaf and check that at least
            ONE meaningful word from `text` actually appears somewhere in the
            combined extracted facts (source_corpus).  If there is zero overlap
            the model invented the text — wipe it and mark the field clearly so
            the UI can display it as missing rather than fabricated data.

            Threshold: at least 1 word of length ≥ 4 characters must appear in
            the corpus (case-insensitive).  Short words (prepositions, articles)
            are skipped to avoid false positives.
            """
            import re
            corpus_lower = source_corpus.lower()

            def has_overlap(text: str) -> bool:
                if not text or not text.strip():
                    return False
                # Skip generic fallback phrases — those are always "valid"
                low = text.lower()
                if ('отсутству' in low or 'нет данных' in low
                        or 'no data' in low or 'not found' in low):
                    return True
                # Long words (≥4 chars)
                words = re.findall(r'[а-яёa-z0-9]{4,}', text.lower())
                if any(w in corpus_lower for w in words):
                    return True
                # Numeric values: dollar amounts, month counts etc. can be short
                # (e.g. "$700" splits at comma → "700" = 3 chars, fails word check)
                nums = re.findall(r'\d+', text)
                return any(n in corpus_lower for n in nums if len(n) >= 2)

            if not isinstance(obj, dict):
                return
            if 'text' in obj and 'source' in obj:
                if not has_overlap(obj.get('text', '')):
                    obj['text'] = 'Данные отсутствуют'
                    obj['source'] = ''
                    obj['has_conflict'] = False
                    obj['conflict_details'] = ''
            for val in obj.values():
                if isinstance(val, dict):
                    _sanitize_hallucinations(val, source_corpus)
                elif isinstance(val, list):
                    for item in val:
                        if isinstance(item, dict):
                            _sanitize_hallucinations(item, source_corpus)

        def _check_cross_field_timeline(doc):
            """
            The LLM only checks timeline/budget within their own field.
            This pass looks for numeric month/week values that appear in
            goals or risks and differ from what timeline field says —
            then flags the timeline field as conflicted if so.
            """
            import re
            if not isinstance(doc, dict):
                return

            def months_in(text):
                """Extract all integer month/week numbers from a string."""
                nums = set()
                for m in re.finditer(r'(\d+)\s*(?:месяц|мес\b|недел)', text, re.IGNORECASE):
                    nums.add(int(m.group(1)))
                return nums

            # Collect all timeline numbers from every field
            all_nums = set()
            for key in ('goals', 'requirements', 'risks'):
                for item in (doc.get(key) or []):
                    text = item.get('text', '') if isinstance(item, dict) else str(item)
                    all_nums |= months_in(text)

            timeline = doc.get('timeline')
            if not isinstance(timeline, dict):
                return
            tl_nums = months_in(timeline.get('text', ''))

            # If timeline field has a number AND other fields have a different number
            conflicting = all_nums - tl_nums
            if tl_nums and conflicting:
                if not timeline.get('has_conflict'):
                    other_vals = ', '.join(str(n) + ' мес.' for n in sorted(conflicting))
                    tl_val = ', '.join(str(n) + ' мес.' for n in sorted(tl_nums))
                    timeline['has_conflict'] = True
                    timeline['conflict_details'] = (
                        f"Конфликт сроков: timeline — {tl_val}; "
                        f"другие поля — {other_vals}"
                    )

        # ── Multi-stage JSON repair ───────────────────────────────────────────
        # LLMs commonly output } where ] is expected (or vice-versa), leave
        # trailing commas, or add markdown fences.  We try progressively more
        # aggressive fixes before giving up.
        def _repair_json(raw: str):
            """Return a parsed dict, or None if all repair stages fail."""
            s = raw.find("{"); e = raw.rfind("}")
            if s == -1 or e == -1:
                return None
            snippet = raw[s:e+1]

            # Stage 1 — direct parse
            try:
                return json.loads(snippet)
            except Exception:
                pass

            # Stage 2 — strip trailing commas: ,} or ,]
            fixed = re.sub(r",\s*([}\]])", r"\1", snippet)
            try:
                return json.loads(fixed)
            except Exception:
                pass

            # Stage 3 — fix bracket mismatches: } where ] expected (or vice-versa)
            # Walk char-by-char tracking the real open-bracket stack so every
            # closer uses whatever bracket matches the opener.  Also closes any
            # brackets that were never closed.
            def _fix_brackets(text: str) -> str:
                out, stack = [], []
                in_str = esc = False
                for ch in text:
                    if esc:
                        esc = False; out.append(ch); continue
                    if ch == "\\" and in_str:
                        esc = True; out.append(ch); continue
                    if ch == '"':
                        in_str = not in_str; out.append(ch); continue
                    if in_str:
                        out.append(ch); continue
                    if ch == "{":
                        stack.append("}"); out.append(ch)
                    elif ch == "[":
                        stack.append("]"); out.append(ch)
                    elif ch in "}]":
                        out.append(stack.pop() if stack else ch)
                    else:
                        out.append(ch)
                while stack:       # close any unclosed brackets
                    out.append(stack.pop())
                return "".join(out)

            fixed2 = _fix_brackets(fixed)
            try:
                return json.loads(fixed2)
            except Exception:
                pass

            # Stage 4 — also strip markdown fences and retry everything
            stripped = re.sub(r"```(?:json)?|```", "", raw).strip()
            for attempt in (stripped, re.sub(r",\s*([}\]])", r"\1", stripped)):
                s2 = attempt.find("{"); e2 = attempt.rfind("}")
                if s2 != -1 and e2 != -1:
                    try:
                        return json.loads(attempt[s2:e2+1])
                    except Exception:
                        try:
                            return json.loads(_fix_brackets(attempt[s2:e2+1]))
                        except Exception:
                            pass

            # Stage 5 — replace single-quoted string values with double-quoted.
            # The model sometimes writes: "text": '{"key": value}' which is
            # invalid JSON.  Replace 'value' → "value" carefully (not inside strings).
            def _fix_single_quotes(text: str) -> str:
                # Replace ': 'value'' → ': "value"' only outside existing double-quoted strings
                return re.sub(r"(?<=:\s)'((?:[^'\\]|\\.)*)'", r'"\1"', text)

            for src_text in (fixed, fixed2, stripped):
                sq = _fix_single_quotes(src_text)
                sq_clean = re.sub(r",\s*([}\]])", r"\1", sq)
                sq_clean = _fix_brackets(sq_clean)
                s3 = sq_clean.find("{"); e3 = sq_clean.rfind("}")
                if s3 != -1 and e3 != -1:
                    try:
                        return json.loads(sq_clean[s3:e3+1])
                    except Exception:
                        pass
            return None

        def _normalize_parsed(data: dict):
            """
            Coerce field types after JSON parse so the rest of the pipeline can
            assume a consistent shape:
              • Scalar fields (timeline, budget, etc.) must be dicts with
                {text, source, has_conflict, conflict_details}.
                If REDUCE returned a plain string or a list, convert it.
              • List fields (goals, requirements, team, risks) must be lists of
                {text, source, …} dicts.  If REDUCE returned a plain list of
                strings, wrap each string.
            """
            if not isinstance(data, dict):
                return
            EMPTY_SCALAR = {"text": "Нет данных", "source": "",
                            "has_conflict": False, "conflict_details": ""}

            for key in SCHEMA_FIELDS_SCALAR:
                val = data.get(key)
                if val is None:
                    data[key] = dict(EMPTY_SCALAR)
                elif isinstance(val, list):
                    # e.g. "timeline": ["6 мес", "12 мес"]  →  join to one string
                    joined = "; ".join(
                        (v.get("text") if isinstance(v, dict) else str(v)).strip()
                        for v in val
                        if (v.get("text") if isinstance(v, dict) else str(v)).strip()
                    )
                    src = next(
                        (v.get("source", "") for v in val if isinstance(v, dict) and v.get("source")),
                        ""
                    )
                    data[key] = {"text": joined or "Нет данных", "source": src,
                                 "has_conflict": False, "conflict_details": ""}
                elif isinstance(val, str):
                    data[key] = {"text": val.strip() or "Нет данных", "source": "",
                                 "has_conflict": False, "conflict_details": ""}
                # else: already a dict — leave it

            for key in SCHEMA_FIELDS_LIST:
                val = data.get(key)
                if val is None:
                    data[key] = []
                elif isinstance(val, str):
                    data[key] = [{"text": val.strip(), "source": "",
                                  "has_conflict": False, "conflict_details": ""}] if val.strip() else []
                elif isinstance(val, list):
                    normalized = []
                    for item in val:
                        if isinstance(item, dict):
                            normalized.append(item)
                        elif isinstance(item, str) and item.strip():
                            normalized.append({"text": item.strip(), "source": "",
                                               "has_conflict": False, "conflict_details": ""})
                    data[key] = normalized
                # else: leave as-is (already a list of dicts)

        # Parse JSON output — now that _repair_json and _merge_partial_specs are defined
        if raw_content is not None:
            # ── Single-REDUCE path ────────────────────────────────────────────
            repaired = _repair_json(raw_content)
            try:
                if repaired is None:
                    raise ValueError("all JSON repair stages failed")
                parsed_data = repaired
                _normalize_parsed(parsed_data)
                _sanitize_placeholders(parsed_data)
                _sanitize_hallucinations(parsed_data, all_extracted_facts)
                _check_cross_field_timeline(parsed_data)
            except Exception:
                parsed_data = {"error": "Llama 3 вернула невалидный JSON", "raw_output": raw_content}
        else:
            # ── Hierarchical-REDUCE path — parse partials and merge ───────────
            parsed_partials = []
            for raw_part in _partial_raw_list:
                p = _repair_json(raw_part)
                if p and isinstance(p, dict):
                    _normalize_parsed(p)
                    _sanitize_placeholders(p)
                    _sanitize_hallucinations(p, all_extracted_facts)
                    parsed_partials.append(p)
            if parsed_partials:
                parsed_data = _merge_partial_specs(parsed_partials)
                _normalize_parsed(parsed_data)
                _check_cross_field_timeline(parsed_data)
            else:
                parsed_data = {"error": "Llama 3 вернула невалидный JSON во всех частях", "raw_output": str(_partial_raw_list)}

        # ── Python-guaranteed post-processing ─────────────────────────────────
        # The REDUCE model sometimes writes "No data" for fields the MAP step
        # successfully extracted.  Fill those gaps directly from `merged` so the
        # final output always reflects everything the parser found.
        # Also: always stamp Python-detected conflict flags — never let the model
        # override what our deterministic comparison already proved.
        if isinstance(parsed_data, dict) and "error" not in parsed_data:

            def _is_empty(val):
                """Return True when a field carries a fallback/empty value."""
                if val is None:
                    return True
                if isinstance(val, dict):
                    t = (val.get("text") or "").strip().lower()
                    return t in ("", "no data", "данные отсутствуют", "нет данных")
                if isinstance(val, list):
                    return len(val) == 0 or all(
                        _is_empty(i) for i in val
                    )
                t = str(val).strip().lower()
                return t in ("", "no data", "данные отсутствуют", "нет данных")

            def _make_fact(entry):
                return {"text": entry["text"], "source": entry["source"],
                        "has_conflict": False, "conflict_details": ""}

            # Scalar fields: inject best merged value when REDUCE said "No data"
            for key in SCHEMA_FIELDS_SCALAR:
                if merged[key] and _is_empty(parsed_data.get(key)):
                    parsed_data[key] = _make_fact(merged[key][0])

            # List fields: inject all merged values when REDUCE said "No data"
            for key in SCHEMA_FIELDS_LIST:
                if merged[key] and _is_empty(parsed_data.get(key)):
                    parsed_data[key] = [_make_fact(e) for e in merged[key]]

            # Strip "Данные отсутствуют" / "Нет данных" placeholder items that
            # sometimes slip into list fields alongside real values
            NON_DATA_LC = {"данные отсутствуют", "нет данных", "no data", ""}
            for key in SCHEMA_FIELDS_LIST:
                items = parsed_data.get(key)
                if isinstance(items, list) and len(items) > 1:
                    cleaned = [i for i in items
                               if not (isinstance(i, dict) and
                                       (i.get("text") or "").strip().lower() in NON_DATA_LC)]
                    if cleaned:
                        parsed_data[key] = cleaned

            # Requirements: remove entries that look like budget section headings
            # (e.g. "Общий бюджет", "Фонд оплаты труда", "Инфраструктура") or
            # architecture section titles ("PostgreSQL Domain Databases",
            # "Event Store Schema") — single-phrase headings with no verb.
            _REQ_HEADING_NOISE = re.compile(
                r"^(?:Общий\s+бюджет|Фонд\s+оплаты\s+труда|Инфраструктура|"
                r"Маркетинг\s+и\s+продажи|Резерв|"
                r"PostgreSQL\s+Domain\s+Databases|Event\s+Store\s+Schema|"
                r"ClickHouse\s+Schema|Data\s+Retention\s+Policy)$",
                re.IGNORECASE,
            )
            req_items = parsed_data.get("requirements")
            if isinstance(req_items, list) and len(req_items) > 1:
                filtered_reqs = [
                    i for i in req_items
                    if not (isinstance(i, dict) and
                            _REQ_HEADING_NOISE.match((i.get("text") or "").strip()))
                ]
                if filtered_reqs:
                    parsed_data["requirements"] = filtered_reqs

            # Architecture / technical_solution: remove values that are just
            # Java package declarations or Python import lists extracted by the
            # code parsers (e.g. "['package io.cryptonest.test']").
            _CODE_NOISE = re.compile(
                r"^\s*\[?'?(?:package|import)\s+[\w\.]+", re.IGNORECASE
            )
            for key in ("architecture", "technical_solution"):
                fld = parsed_data.get(key)
                if isinstance(fld, dict):
                    if _CODE_NOISE.match(fld.get("text") or ""):
                        # Replace with best merged value if available
                        if merged[key]:
                            parsed_data[key] = _make_fact(merged[key][0])
                        else:
                            fld["text"] = "Нет данных"

            # Team: remove entries that look like spec/API descriptions rather
            # than actual people, and pure numbers leaked from config.json.
            _SPEC_NOISE = re.compile(
                r"^(?:GET|POST|PUT|DELETE|PATCH)\s+/|https?://|"
                r"^(?:RESTful|OAuth|AES|HL7|FHIR|шифрован|аудит.лог|"
                # Technical component names that appear in architecture diagrams
                r"Kong\s+API\s+Gateway|Business\s+Domain\s+Services|"
                r"AI/ML\s+Services|Background\s+Workers|WebSocket\s+Hub|"
                r"OIDC\s+Provider|Auth\s+Service|RBAC\s+Model|"
                r"Notification\s+Hub|Event\s+Bus)",
                re.IGNORECASE,
            )
            team_items = parsed_data.get("team")
            if isinstance(team_items, list) and len(team_items) > 1:
                cleaned_team = []
                for i in team_items:
                    if not isinstance(i, dict):
                        continue
                    txt = (i.get("text") or "").strip()
                    # Skip spec/API noise
                    if _SPEC_NOISE.search(txt):
                        continue
                    # Skip pure-number entries (e.g. "21", "52" from config.json values)
                    if re.match(r"^\d+$", txt):
                        continue
                    cleaned_team.append(i)
                if cleaned_team:
                    parsed_data["team"] = cleaned_team

            # Budget text cleanup: if the model stored a Python-list repr or
            # JSON array string (e.g. "[{'amount': '30 000 000 тенге', ...}]"),
            # extract the first monetary amount with a simple regex.
            b = parsed_data.get("budget")
            if isinstance(b, dict):
                bt = (b.get("text") or "").strip()
                if bt.startswith("[") or bt.startswith("{") or bt.startswith("'"):
                    # Try to pull the first number+currency from the string
                    m = re.search(r"[\d\s]+(?:тенге|тнг|руб|usd|\$|€|₸)", bt, re.IGNORECASE)
                    if m:
                        b["text"] = m.group(0).strip()
                    elif merged["budget"]:
                        b["text"]   = merged["budget"][0]["text"]
                        b["source"] = merged["budget"][0]["source"]

            # Conflict flags: always trust Python detection over the model.
            # Also fix the common mistake where the model writes the conflict
            # description into the "text" field instead of the actual value.
            if budget_conflict and merged["budget"]:
                b = parsed_data.get("budget")
                if isinstance(b, dict):
                    # If text looks like a conflict description, replace with first real value
                    if ("conflict" in (b.get("text") or "").lower()
                            or _is_empty(b)):
                        b["text"]   = merged["budget"][0]["text"]
                        b["source"] = merged["budget"][0]["source"]
                    b["has_conflict"]     = True
                    b["conflict_details"] = budget_detail

            if timeline_conflict and merged["timeline"]:
                t = parsed_data.get("timeline")
                if isinstance(t, dict):
                    if ("conflict" in (t.get("text") or "").lower()
                            or _is_empty(t)):
                        t["text"]   = merged["timeline"][0]["text"]
                        t["source"] = merged["timeline"][0]["source"]
                    t["has_conflict"]     = True
                    t["conflict_details"] = timeline_detail

        def count_conflicts(doc):
            count = 0
            if not isinstance(doc, dict): return 0
            for value in doc.values():
                if isinstance(value, dict) and value.get("has_conflict"):
                    count += 1
                elif isinstance(value, list):
                    for item in value:
                        if isinstance(item, dict) and item.get("has_conflict"):
                            count += 1
            return count

        def _to_spec_document(rich):
            """
            Convert the rich {text, source, has_conflict, conflict_details} format
            into a spec-compliant document where scalar fields are plain strings
            and array fields are plain string lists.
            """
            def txt(f):
                if isinstance(f, dict):
                    return (f.get('text') or '').strip()
                if isinstance(f, list):
                    # REDUCE occasionally writes a list for a scalar field
                    parts = [txt(i) for i in f if txt(i)]
                    return "; ".join(parts)
                return str(f or '').strip()

            def lst(arr):
                if not isinstance(arr, list):
                    return []
                return [txt(i) for i in arr if i]

            return {
                "project_overview":   rich.get("project_overview", ""),
                "goals":              lst(rich.get("goals")),
                "requirements":       lst(rich.get("requirements")),
                "technical_solution": txt(rich.get("technical_solution")),
                "architecture":       txt(rich.get("architecture")),
                "team":               lst(rich.get("team")),
                "timeline":           txt(rich.get("timeline")),
                "budget":             txt(rich.get("budget")),
                "risks":              lst(rich.get("risks")),
            }

        # Build both the spec-compliant document and the rich extended document
        spec_document = _to_spec_document(parsed_data) if "error" not in parsed_data else parsed_data

        return {
            "document":          spec_document,    # spec-compliant plain strings/lists
            "document_extended": parsed_data,       # rich {text,source,has_conflict,...} for UI
            "metadata": {
                "model_name": f"Local GPU Map-Reduce ({self.model_name})",
                "llm_calls": len(files_data) + 1,
                "total_tokens": total_tokens_used,
                "duration_ms": duration_ms,
                "conflicts_found": count_conflicts(parsed_data)
            },
            "trace": {
                "steps": [
                    "Парсинг файлов",
                    f"Map: Извлечение фактов (обработано файлов: {len(files_data)})",
                    "Reduce: Сборка JSON"
                ]
            }
        }

# ─────────────────────────────────────────────────────────────────────────────
# FASTAPI APP
# ─────────────────────────────────────────────────────────────────────────────
app = FastAPI()
# When running in Docker, LM Studio is on the host machine.
# Set LM_STUDIO_HOST env var to override (e.g. "host.docker.internal" on Docker Desktop).
LOCAL_PC_IP  = os.environ.get("LM_STUDIO_HOST",  "host.docker.internal")
# Set LM_STUDIO_MODEL to match whatever model is loaded in LM Studio.
# LM Studio ignores the model field and always uses the loaded model,
# but the value appears in the output metadata so keep it accurate.
LM_MODEL     = os.environ.get("LM_STUDIO_MODEL", "qwen/qwen2.5-v1-7b")
current_llm  = LocalProvider(base_url=f"http://{LOCAL_PC_IP}:1234/v1", model_name=LM_MODEL)

# ── Archive API ───────────────────────────────────────────────────────────────
@app.get("/api/samples")
async def list_samples():
    """Return sorted list of available pre-generated JSON samples."""
    if not os.path.isdir(SAMPLES_DIR):
        return JSONResponse({"files": [], "dir": SAMPLES_DIR})
    files = sorted(f for f in os.listdir(SAMPLES_DIR) if f.endswith(".json"))
    return JSONResponse({"files": files, "count": len(files)})

@app.get("/api/samples/{filename}")
async def get_sample(filename: str):
    """Return the content of one sample JSON."""
    # Security: only allow simple filenames, no path traversal
    if "/" in filename or "\\" in filename or ".." in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    path = os.path.join(SAMPLES_DIR, filename)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="File not found")
    with open(path, encoding="utf-8") as f:
        return JSONResponse(json.load(f))

# ── Main page ─────────────────────────────────────────────────────────────────
@app.get("/")
async def main_page():
    html_content = r"""
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Генератор проектной документации</title>
    <style id="theme-styles">
        :root {
            --primary: #4F46E5; --primary-hover: #4338CA;
            --bg: #F3F4F6; --card-bg: #FFFFFF;
            --text-main: #1F2937; --text-muted: #6B7280;
            --border: #E5E7EB;
            --source-bg: #DBEAFE; --source-text: #1E40AF;
            --green: #10B981; --green-hover: #059669;
            --conflict-bg: #FEF3C7; --conflict-text: #92400E;
            --shadow: 0 10px 15px -3px rgba(0,0,0,.1);
        }
        * { box-sizing: border-box; }
        body { font-family: 'Segoe UI', system-ui, -apple-system, sans-serif; background: var(--bg); color: var(--text-main); padding: 40px 20px; margin: 0; line-height: 1.5; }
        .container { max-width: 800px; margin: 0 auto; background: var(--card-bg); padding: 40px; border-radius: 12px; box-shadow: 0 10px 15px -3px rgba(0,0,0,.1); }
        h1, h2, h3 { color: var(--text-main); }
        h1 { text-align: center; margin-top: 0; }
        p { color: var(--text-muted); }

        /* ── Tabs ── */
        .tab-bar { display: flex; gap: 8px; margin-bottom: 28px; border-bottom: 2px solid var(--border); padding-bottom: 0; }
        .tab-btn { background: none; border: none; padding: 10px 20px; font-size: 15px; font-weight: 600; color: var(--text-muted); cursor: pointer; border-bottom: 3px solid transparent; margin-bottom: -2px; border-radius: 0; width: auto; transition: color .2s, border-color .2s; }
        .tab-btn:hover { color: var(--primary); }
        .tab-btn.active { color: var(--primary); border-bottom-color: var(--primary); }

        /* ── Upload zone ── */
        .upload-zone { border: 2px dashed var(--border); padding: 32px; text-align: center; border-radius: 8px; margin-bottom: 20px; transition: border-color .3s; }
        .upload-zone:hover { border-color: var(--primary); }
        input[type="file"] { margin-bottom: 10px; }

        /* ── Buttons ── */
        button { background: var(--primary); color: white; padding: 12px 24px; border: none; border-radius: 6px; font-size: 16px; font-weight: 600; cursor: pointer; width: 100%; transition: background .2s; }
        button:hover { background: var(--primary-hover); }
        .btn-green { background: var(--green); }
        .btn-green:hover { background: var(--green-hover); }
        .btn-gray { background: var(--text-muted); }
        .btn-gray:hover { background: #4B5563; }
        .btn-row { display: flex; gap: 10px; margin-top: 10px; }
        .btn-row button { flex: 1; }

        /* ── Loading ── */
        #loading-screen { display: none; text-align: center; padding: 40px 0; }
        .spinner { width: 60px; height: 60px; border: 5px solid var(--border); border-top: 5px solid var(--primary); border-radius: 50%; animation: spin 1s linear infinite; margin: 0 auto 20px; }
        @keyframes spin { to { transform: rotate(360deg); } }

        /* ── Progress Visualization ── */
        .progress-container { max-width: 600px; margin: 30px auto; text-align: left; }
        .phase-header { display: flex; align-items: center; gap: 10px; margin-bottom: 15px; font-weight: 600; color: var(--text-main); }
        .phase-badge { background: var(--primary); color: white; padding: 4px 12px; border-radius: 20px; font-size: 12px; }
        .file-progress-list { display: flex; flex-direction: column; gap: 8px; }
        .file-progress-item { display: flex; align-items: center; gap: 12px; padding: 10px 15px; background: var(--bg); border-radius: 8px; border: 1px solid var(--border); transition: all .3s; }
        .file-progress-item.processing { border-color: var(--primary); background: var(--source-bg); }
        .file-progress-item.done { border-color: var(--green); }
        .file-icon { font-size: 20px; }
        .file-name { flex: 1; font-size: 14px; color: var(--text-main); }
        .file-status { font-size: 12px; color: var(--text-muted); }
        .mini-spinner { width: 16px; height: 16px; border: 2px solid var(--border); border-top-color: var(--primary); border-radius: 50%; animation: spin 1s linear infinite; }
        .checkmark { color: var(--green); font-size: 18px; font-weight: bold; }
        .overall-progress { margin-top: 25px; }
        .progress-bar-bg { height: 8px; background: var(--border); border-radius: 4px; overflow: hidden; }
        .progress-bar-fill { height: 100%; background: linear-gradient(90deg, var(--primary), var(--green)); width: 0%; transition: width .5s ease; border-radius: 4px; }
        .progress-stats { display: flex; justify-content: space-between; margin-top: 10px; font-size: 13px; color: var(--text-muted); }

        /* ── Result ── */
        #result-screen { display: none; }
        .result-section { margin-bottom: 24px; padding-bottom: 16px; border-bottom: 1px solid var(--border); }
        .result-section h3 { color: var(--primary); margin-bottom: 8px; font-size: 18px; }
        ul { margin-top: 0; padding-left: 20px; }
        li { margin-bottom: 10px; color: var(--text-main); }
        .source-badge { display: inline-block; background: var(--source-bg); color: var(--source-text); font-size: 11px; font-weight: 600; padding: 2px 8px; border-radius: 12px; margin-left: 8px; vertical-align: middle; }
        .meta-data { background: var(--bg); padding: 15px; border-radius: 6px; font-size: 14px; margin-top: 20px; border-left: 4px solid var(--primary); }

        /* ── Archive browser ── */
        .archive-header { display: flex; align-items: center; justify-content: space-between; margin-bottom: 12px; }
        .archive-header h3 { margin: 0; font-size: 16px; }
        .archive-count { font-size: 13px; color: var(--text-muted); background: var(--bg); padding: 3px 10px; border-radius: 10px; }
        .archive-search { width: 100%; padding: 9px 12px; border: 1px solid var(--border); border-radius: 6px; font-size: 14px; margin-bottom: 12px; outline: none; }
        .archive-search:focus { border-color: var(--primary); box-shadow: 0 0 0 3px rgba(79,70,229,.15); }
        .archive-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(220px, 1fr)); gap: 10px; max-height: 380px; overflow-y: auto; padding-right: 4px; }
        .archive-card { border: 1px solid var(--border); border-radius: 8px; padding: 12px 14px; cursor: pointer; transition: border-color .2s, box-shadow .2s, background .2s; }
        .archive-card:hover { border-color: var(--primary); box-shadow: 0 2px 8px rgba(79,70,229,.15); background: #F5F3FF; }
        .archive-card .card-name { font-weight: 600; font-size: 13px; color: var(--text-main); }
        .archive-card .card-meta { font-size: 12px; color: var(--text-muted); margin-top: 4px; }
        .archive-card .card-conflicts { display: inline-block; font-size: 11px; font-weight: 600; margin-top: 6px; padding: 2px 8px; border-radius: 10px; background: #FEF3C7; color: #92400E; }
        .archive-card .card-conflicts.none { background: #D1FAE5; color: #065F46; }
        .archive-empty { text-align: center; padding: 40px 20px; color: var(--text-muted); font-size: 14px; }
        .divider { display: flex; align-items: center; gap: 12px; margin: 20px 0; color: var(--text-muted); font-size: 13px; }
        .divider::before, .divider::after { content: ''; flex: 1; height: 1px; background: var(--border); }

        /* ── JSON viewer source badge (viewer mode indicator) ── */
        .viewer-badge { display: inline-block; background: #EDE9FE; color: #6D28D9; font-size: 12px; font-weight: 600; padding: 3px 10px; border-radius: 8px; margin-left: 10px; vertical-align: middle; }

        /* ── Completeness Score ── */
        .score-badge { display: inline-flex; align-items: center; gap: 6px; background: linear-gradient(135deg, var(--primary), var(--green)); color: white; padding: 6px 14px; border-radius: 20px; font-size: 14px; font-weight: 700; margin-left: 10px; vertical-align: middle; }
        .score-badge.low { background: linear-gradient(135deg, #EF4444, #F59E0B); }
        .score-badge.medium { background: linear-gradient(135deg, #F59E0B, #10B981); }

        /* ── Conflict Modal ── */
        .modal-overlay { display: none; position: fixed; top: 0; left: 0; right: 0; bottom: 0; background: rgba(0,0,0,0.6); backdrop-filter: blur(4px); z-index: 2000; justify-content: center; align-items: center; padding: 20px; }
        .modal-content { background: var(--card-bg); border-radius: 16px; max-width: 700px; width: 100%; max-height: 80vh; overflow: hidden; box-shadow: 0 25px 50px -12px rgba(0,0,0,0.25); }
        .modal-header { display: flex; align-items: center; justify-content: space-between; padding: 20px 24px; border-bottom: 1px solid var(--border); }
        .modal-header h3 { margin: 0; color: var(--conflict-text); }
        .modal-close { background: none; border: none; font-size: 24px; color: var(--text-muted); cursor: pointer; width: auto; padding: 0; }
        .modal-body { padding: 24px; overflow-y: auto; max-height: 60vh; }
        .diff-container { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; margin-bottom: 20px; }
        .diff-box { background: var(--bg); border-radius: 8px; padding: 16px; border: 2px solid var(--border); }
        .diff-box.highlight { border-color: var(--conflict-text); background: var(--conflict-bg); }
        .diff-box h4 { margin: 0 0 10px; font-size: 13px; color: var(--text-muted); text-transform: uppercase; }
        .diff-box p { margin: 0; font-size: 14px; color: var(--text-main); }
        .conflict-actions { display: flex; gap: 10px; justify-content: center; margin-top: 20px; }
        .conflict-actions button { width: auto; padding: 10px 20px; }

        /* ── Enhanced Result Cards ── */
        .result-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; }
        .result-card { background: var(--bg); border-radius: 12px; padding: 20px; border: 1px solid var(--border); transition: transform .2s, box-shadow .2s; }
        .result-card:hover { transform: translateY(-2px); box-shadow: var(--shadow); }
        .result-card h4 { margin: 0 0 12px; color: var(--primary); font-size: 14px; text-transform: uppercase; }
        .result-card .value { font-size: 24px; font-weight: 700; color: var(--text-main); }
        .result-card .label { font-size: 12px; color: var(--text-muted); margin-top: 4px; }

        /* ── Animations ── */
        @keyframes fadeIn { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; transform: translateY(0); } }
        .fade-in { animation: fadeIn 0.4s ease; }
        @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
        .pulse { animation: pulse 2s ease-in-out infinite; }
    </style>
</head>
<body>
<div class="container">
    <h1>
        🧠 NeuralDocs
        <span style="font-size:12px;background:var(--primary);color:white;padding:4px 10px;border-radius:20px;vertical-align:middle;margin-left:8px;">AI</span>
    </h1>
    <p style="text-align:center;margin-top:-10px;margin-bottom:25px;color:var(--text-muted);">Интеллектуальный анализатор проектной документации</p>

    <!-- ═══════════════════════ UPLOAD SCREEN ═══════════════════════ -->
    <div id="upload-screen">
        <div class="tab-bar">
            <button class="tab-btn active" onclick="switchTab('generate')">📝 Создать документ</button>
            <button class="tab-btn" onclick="switchTab('load')">📂 Загрузить документ</button>
        </div>

        <!-- Tab 1: Generate -->
        <div id="tab-generate">
            <p style="text-align:center;margin-bottom:20px;">Загрузите файлы проекта — мы извлечем структурированную информацию и найдем противоречия между документами.</p>
            <form id="upload-form">
                <div class="upload-zone">
                    <p style="margin:0 0 15px;color:var(--text-muted);">📁 Перетащите файлы или выберите</p>
                    <input type="file" id="file-input" name="files" multiple required accept=".txt,.md,.json,.py,.java,.docx,.pdf">
                    <p style="margin:10px 0 0;font-size:12px;color:var(--text-muted);">Поддерживаемые форматы: TXT, MD, JSON, PY, Java, DOCX, PDF</p>
                </div>
                <button type="submit">🔍 Анализировать документы</button>
            </form>
        </div>

        <!-- Tab 2: Load JSON -->
        <div id="tab-load" style="display:none;">

            <!-- Local file picker -->
            <div class="upload-zone" id="json-drop-zone">
                <p style="margin:0 0 10px;font-weight:600;color:var(--text-main);">📤 Загрузить JSON с устройства</p>
                <input type="file" id="json-file-input" accept=".json">
                <p style="margin:8px 0 0;font-size:13px;color:var(--text-muted);">Откройте ранее сохраненный документ</p>
            </div>

            <div class="divider">или</div>

            <!-- Archive browser -->
            <div class="archive-header">
                <h3>📚 Библиотека примеров</h3>
                <span class="archive-count" id="archive-count">загрузка...</span>
            </div>
            <input class="archive-search" type="text" id="archive-search" placeholder="🔍 Поиск по проектам..." oninput="filterArchive()">
            <div class="archive-grid" id="archive-grid">
                <div class="archive-empty">Загрузка библиотеки...</div>
            </div>
        </div>
    </div>

    <!-- ═══════════════════════ LOADING SCREEN ═══════════════════════ -->
    <div id="loading-screen">
        <div class="spinner"></div>
        <h3>Анализ документов...</h3>
        <p style="color:var(--text-muted);">Нейросеть извлекает информацию из каждого файла</p>

        <div class="progress-container">
            <div class="phase-header">
                <span class="phase-badge">ШАГ 1</span>
                <span>Чтение и анализ файлов</span>
            </div>
            <div class="file-progress-list" id="file-progress-list">
                <!-- Dynamically populated -->
            </div>

            <div class="phase-header" style="margin-top: 20px;">
                <span class="phase-badge" style="background: var(--green);">ШАГ 2</span>
                <span>Сборка итогового документа</span>
            </div>
            <div class="overall-progress">
                <div class="progress-bar-bg">
                    <div class="progress-bar-fill" id="overall-progress-bar"></div>
                </div>
                <div class="progress-stats">
                    <span id="progress-text">Обработано: 0 файлов</span>
                    <span id="token-counter">0 токенов</span>
                </div>
            </div>
        </div>
    </div>

    <!-- ═══════════════════════ RESULT SCREEN ═══════════════════════ -->
    <div id="result-screen">
        <h2 id="result-title">📋 Структурированный документ</h2>
        <div id="parsed-content"></div>
        <div class="meta-data" id="meta-content"></div>
        <div class="btn-row" style="margin-top:20px;">
            <button class="btn-green" id="download-btn">📥 Скачать JSON</button>
            <button class="btn-gray" id="load-another-btn" style="display:none;">📄 Загрузить другой JSON</button>
        </div>
        <button class="btn-gray" id="reset-btn" style="margin-top:10px;">← Начать заново</button>
    </div>
</div>

<!-- ═══════════════════════ CONFLICT MODAL ═══════════════════════ -->
<div class="modal-overlay" id="conflict-modal" onclick="closeModalOnOverlay(event)">
    <div class="modal-content">
        <div class="modal-header">
            <h3>⚠️ Конфликт данных</h3>
            <button class="modal-close" onclick="closeConflictModal()">&times;</button>
        </div>
        <div class="modal-body">
            <p style="color: var(--text-muted); margin-bottom: 20px;">Обнаружены различия в данных из разных источников:</p>
            <div class="diff-container" id="diff-container">
                <!-- Dynamically populated -->
            </div>
            <p style="text-align: center; color: var(--text-muted); font-size: 13px;">Система автоматически выбрала последнее значение. Проверьте исходные документы.</p>
        </div>
    </div>
</div>

<script>
// ─────────────────────────────────────────────────────────────────────────────
// STATE
// ─────────────────────────────────────────────────────────────────────────────
let currentRawData = null;
let allArchiveFiles = [];   // [{name, conflicts, tokens, duration}]
let viewerMode = false;     // true when viewing a loaded JSON (not generated)

// ─────────────────────────────────────────────────────────────────────────────
// PROGRESS VISUALIZATION
// ─────────────────────────────────────────────────────────────────────────────
let fileProgressState = [];

function initFileProgress(fileNames) {
    fileProgressState = fileNames.map(name => ({ name, status: 'pending', tokens: 0 }));
    renderFileProgress();
}

function updateFileProgress(fileName, status, tokens = 0) {
    const file = fileProgressState.find(f => f.name === fileName);
    if (file) {
        file.status = status;
        file.tokens = tokens;
    }
    renderFileProgress();
    updateOverallProgress();
}

function renderFileProgress() {
    const container = document.getElementById('file-progress-list');
    container.innerHTML = fileProgressState.map(f => {
        const icon = f.status === 'done' ? '<span class="checkmark">✓</span>' :
                     f.status === 'processing' ? '<div class="mini-spinner"></div>' :
                     '<span style="color: var(--text-muted);">○</span>';
        const cls = f.status === 'processing' ? 'processing' : f.status === 'done' ? 'done' : '';
        const statusText = f.status === 'done' ? `${f.tokens} токенов` :
                           f.status === 'processing' ? 'Обработка...' : 'Ожидание';
        return `
            <div class="file-progress-item ${cls}">
                <span class="file-icon">📄</span>
                <span class="file-name">${f.name}</span>
                <span class="file-status">${icon} ${statusText}</span>
            </div>
        `;
    }).join('');
}

function updateOverallProgress() {
    const total = fileProgressState.length;
    const done = fileProgressState.filter(f => f.status === 'done').length;
    const percent = total > 0 ? (done / total) * 100 : 0;
    document.getElementById('overall-progress-bar').style.width = percent + '%';
    document.getElementById('progress-text').textContent = `Файлов обработано: ${done}/${total}`;
    const totalTokens = fileProgressState.reduce((sum, f) => sum + f.tokens, 0);
    document.getElementById('token-counter').textContent = totalTokens.toLocaleString() + ' токенов';
}

// ─────────────────────────────────────────────────────────────────────────────
// CONFLICT MODAL
// ─────────────────────────────────────────────────────────────────────────────
let currentConflict = null;

function showConflictModal(fieldName, conflictDetails, sources) {
    const modal = document.getElementById('conflict-modal');
    const container = document.getElementById('diff-container');

    // Parse conflict details to extract sources
    const diffHtml = sources.map((src, idx) => `
        <div class="diff-box ${idx === sources.length - 1 ? 'highlight' : ''}">
            <h4>${src.file}, строка ${src.line}</h4>
            <p>${src.value}</p>
        </div>
    `).join('');

    container.innerHTML = diffHtml;
    modal.style.display = 'flex';
}

function closeConflictModal() {
    document.getElementById('conflict-modal').style.display = 'none';
}

function closeModalOnOverlay(e) {
    if (e.target === e.currentTarget) closeConflictModal();
}

// ─────────────────────────────────────────────────────────────────────────────
// COMPLETENESS SCORE
// ─────────────────────────────────────────────────────────────────────────────
function calculateCompletenessScore(doc) {
    const fields = [
        { key: 'project_overview', weight: 10 },
        { key: 'goals', weight: 15, isArray: true },
        { key: 'requirements', weight: 15, isArray: true },
        { key: 'technical_solution', weight: 15 },
        { key: 'architecture', weight: 10 },
        { key: 'team', weight: 10, isArray: true },
        { key: 'timeline', weight: 10 },
        { key: 'budget', weight: 10 },
        { key: 'risks', weight: 5, isArray: true }
    ];

    let score = 0;
    fields.forEach(field => {
        const value = doc[field.key];
        if (field.isArray) {
            if (value && value.length > 0) {
                // Rich format: [{text: ...}]  |  Spec format: ["string", ...]
                const first = value[0];
                const text = (typeof first === 'object' && first !== null) ? (first.text || '') : String(first || '');
                if (text && !text.includes('отсутству')) score += field.weight;
            }
        } else {
            // Rich format: {text: ...}  |  Spec format: "string"
            const text = (value && typeof value === 'object') ? (value.text || '') : String(value || '');
            if (text && text.length > 10 && !text.includes('отсутству')) score += field.weight;
        }
    });

    return Math.round(score);
}

function getScoreClass(score) {
    if (score >= 80) return '';
    if (score >= 50) return 'medium';
    return 'low';
}

// ─────────────────────────────────────────────────────────────────────────────
// TAB SWITCHING
// ─────────────────────────────────────────────────────────────────────────────
function switchTab(tab) {
    document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
    document.getElementById('tab-generate').style.display = 'none';
    document.getElementById('tab-load').style.display = 'none';

    if (tab === 'generate') {
        document.getElementById('tab-generate').style.display = 'block';
        document.querySelectorAll('.tab-btn')[0].classList.add('active');
    } else {
        document.getElementById('tab-load').style.display = 'block';
        document.querySelectorAll('.tab-btn')[1].classList.add('active');
        loadArchiveList();
    }
}

// ─────────────────────────────────────────────────────────────────────────────
// GENERATE FORM
// ─────────────────────────────────────────────────────────────────────────────
document.getElementById('upload-form').addEventListener('submit', async (e) => {
    e.preventDefault();
    viewerMode = false;

    const files = document.getElementById('file-input').files;
    const fileNames = Array.from(files).map(f => f.name);

    // Initialize progress visualization
    initFileProgress(fileNames);
    showScreen('loading');

    // Simulate progress for each file (since we don't have streaming yet)
    const progressInterval = setInterval(() => {
        const pendingFiles = fileProgressState.filter(f => f.status === 'pending');
        if (pendingFiles.length > 0) {
            const nextFile = pendingFiles[0];
            updateFileProgress(nextFile.name, 'processing');
            setTimeout(() => {
                updateFileProgress(nextFile.name, 'done', Math.floor(Math.random() * 500) + 200);
            }, 800 + Math.random() * 1000);
        }
    }, 1500);

    const formData = new FormData();
    for (let i = 0; i < files.length; i++) formData.append('files', files[i]);
    
    // Default language is Russian
    formData.append('language', 'ru');

    try {
        const resp = await fetch('/generate_document', { method: 'POST', body: formData });
        clearInterval(progressInterval);
        if (!resp.ok) {
            const err = await resp.json();
            throw new Error(err.detail || 'Неизвестная ошибка сервера');
        }
        const data = await resp.json();
        showResult(data, false);
    } catch (err) {
        clearInterval(progressInterval);
        alert('Ошибка: ' + err.message);
        showScreen('upload');
    }
});

// ─────────────────────────────────────────────────────────────────────────────
// LOCAL JSON FILE PICKER
// ─────────────────────────────────────────────────────────────────────────────
document.getElementById('json-file-input').addEventListener('change', (e) => {
    const file = e.target.files[0];
    if (!file) return;
    const reader = new FileReader();
    reader.onload = (ev) => {
        try {
            const data = JSON.parse(ev.target.result);
            showResult(data, true, file.name);
        } catch {
            alert('Ошибка: файл не является валидным JSON');
        }
    };
    reader.readAsText(file, 'utf-8');
});

// ─────────────────────────────────────────────────────────────────────────────
// ARCHIVE BROWSER
// ─────────────────────────────────────────────────────────────────────────────
async function loadArchiveList() {
    if (allArchiveFiles.length > 0) return; // already loaded
    const grid = document.getElementById('archive-grid');
    try {
        const resp = await fetch('/api/samples');
        const data = await resp.json();
        document.getElementById('archive-count').textContent = `${data.count} файлов`;

        if (!data.files || data.files.length === 0) {
            grid.innerHTML = '<div class="archive-empty">Архив пуст. Запустите setup_v2.py для генерации образцов.</div>';
            return;
        }

        // Load metadata for each file (lightweight — just parse conflicts from filename hint via API)
        // We fetch all in parallel with Promise.all but limit batch to avoid flooding
        const BATCH = 20;
        allArchiveFiles = [];
        for (let i = 0; i < data.files.length; i += BATCH) {
            const batch = data.files.slice(i, i + BATCH);
            const results = await Promise.all(batch.map(name =>
                fetch(`/api/samples/${name}`)
                    .then(r => r.json())
                    .then(d => ({
                        name,
                        overview: (d.document && d.document.project_overview) ? d.document.project_overview.slice(0, 60) : '',
                        conflicts: (d.metadata && d.metadata.conflicts_found) || 0,
                        tokens: (d.metadata && d.metadata.total_tokens) || 0,
                        duration: (d.metadata && d.metadata.duration_ms) || 0,
                    }))
                    .catch(() => ({ name, overview: '', conflicts: 0, tokens: 0, duration: 0 }))
            ));
            allArchiveFiles.push(...results);
        }
        renderArchiveGrid(allArchiveFiles);
    } catch (err) {
        grid.innerHTML = `<div class="archive-empty">Не удалось загрузить архив: ${err.message}</div>`;
    }
}

function renderArchiveGrid(files) {
    const grid = document.getElementById('archive-grid');
    if (!files.length) {
        grid.innerHTML = '<div class="archive-empty">Ничего не найдено</div>';
        return;
    }
    grid.innerHTML = files.map(f => `
        <div class="archive-card" onclick="loadSampleByName('${f.name}')">
            <div class="card-name">📄 ${f.name}</div>
            <div class="card-meta" title="${f.overview}">${f.overview || 'Нет описания'}...</div>
            <div class="card-meta">🪙 ${f.tokens} токенов · ⏱ ${f.duration} мс</div>
            <span class="card-conflicts ${f.conflicts === 0 ? 'none' : ''}">
                ${f.conflicts === 0 ? '✓ Конфликтов нет' : '⚠️ Конфликтов: ' + f.conflicts}
            </span>
        </div>
    `).join('');
}

function filterArchive() {
    const q = document.getElementById('archive-search').value.toLowerCase();
    const filtered = allArchiveFiles.filter(f =>
        f.name.toLowerCase().includes(q) || f.overview.toLowerCase().includes(q)
    );
    renderArchiveGrid(filtered);
}

async function loadSampleByName(filename) {
    showScreen('loading');
    try {
        const resp = await fetch(`/api/samples/${encodeURIComponent(filename)}`);
        if (!resp.ok) throw new Error('Файл не найден');
        const data = await resp.json();
        showResult(data, true, filename);
    } catch (err) {
        alert('Ошибка загрузки: ' + err.message);
        showScreen('upload');
    }
}

// ─────────────────────────────────────────────────────────────────────────────
// SHOW RESULT
// ─────────────────────────────────────────────────────────────────────────────
function showResult(data, isViewer, filename) {
    viewerMode = isViewer;
    currentRawData = data;

    const title = document.getElementById('result-title');
    const loadAnotherBtn = document.getElementById('load-another-btn');
    // Use the rich extended document for UI rendering if available (spec-compliant
    // plain-string `document` is for download/fine-tuning only)
    const doc = data.document_extended || data.document || {};

    // Calculate and show completeness score
    const score = calculateCompletenessScore(doc);
    const scoreClass = getScoreClass(score);
    const scoreEmoji = score >= 80 ? '✅' : score >= 50 ? '⚠️' : '❌';

    if (isViewer) {
        title.innerHTML = `Просмотр документа <span class="viewer-badge">📄 ${filename || 'JSON'}</span><span class="score-badge ${scoreClass}">${scoreEmoji} ${score}%</span>`;
        loadAnotherBtn.style.display = 'block';
    } else {
        title.innerHTML = `Результат генерации: <span class="score-badge ${scoreClass}">${scoreEmoji} ${score}%</span>`;
        loadAnotherBtn.style.display = 'none';
    }

    renderResults(data);
    showScreen('result');
}

// ─────────────────────────────────────────────────────────────────────────────
// SCREEN SWITCHER
// ─────────────────────────────────────────────────────────────────────────────
function showScreen(name) {
    document.getElementById('upload-screen').style.display  = name === 'upload'  ? 'block' : 'none';
    document.getElementById('loading-screen').style.display = name === 'loading' ? 'block' : 'none';
    document.getElementById('result-screen').style.display  = name === 'result'  ? 'block' : 'none';
}

// ─────────────────────────────────────────────────────────────────────────────
// RENDER RESULTS  (unchanged logic — used by both generate and viewer paths)
// ─────────────────────────────────────────────────────────────────────────────
function renderResults(data) {
    // Prefer rich extended document for UI (has source badges & conflict details).
    // Plain-string spec `document` is used for download/fine-tuning.
    const doc  = data.document_extended || data.document || {};
    const meta = data.metadata  || {};
    const container = document.getElementById('parsed-content');
    
    // Get current language from dropdown
    const currentLang = document.getElementById('output-language')?.value || 'ru';
    
    // Multilingual labels
    const labels = {
        ru: {
            project_overview: 'Обзор проекта',
            goals: 'Цели проекта',
            requirements: 'Требования',
            technical_solution: 'Техническое решение',
            architecture: 'Архитектура',
            team: 'Команда',
            timeline: 'Сроки',
            budget: 'Бюджет',
            risks: 'Риски',
            no_data: 'Нет данных',
            conflict_click: 'Конфликт (клик для деталей)',
            conflicts_found: 'Обнаружено противоречий',
            conflict_desc: 'Найдены несовместимые данные в документах:',
            metadata: 'Статистика обработки',
            files_processed: 'Файлов проанализировано',
            final_assembly: 'Сборка документа',
            tokens: 'Токенов обработано',
            duration: 'Время обработки',
            ms: 'мс'
        },
        en: {
            project_overview: 'Project Overview',
            goals: 'Project Goals',
            requirements: 'Requirements',
            technical_solution: 'Technical Solution',
            architecture: 'Architecture',
            team: 'Team',
            timeline: 'Timeline',
            budget: 'Budget',
            risks: 'Risks',
            no_data: 'No data available',
            conflict_click: 'Conflict (click for details)',
            conflicts_found: 'Conflicts detected',
            conflict_desc: 'Inconsistent data found across documents:',
            metadata: 'Processing statistics',
            files_processed: 'Files analyzed',
            final_assembly: 'Document assembly',
            tokens: 'Tokens processed',
            duration: 'Processing time',
            ms: 'ms'
        },
        kz: {
            project_overview: 'Жоба шолуы',
            goals: 'Жоба мақсаттары',
            requirements: 'Талаптар',
            technical_solution: 'Техникалық шешім',
            architecture: 'Архитектура',
            team: 'Команда',
            timeline: 'Мерзімдер',
            budget: 'Бюджет',
            risks: 'Тәуекелдер',
            no_data: 'Деректер жоқ',
            conflict_click: 'Қақтығыс (толықтыру үшін басыңыз)',
            conflicts_found: 'Қақтығыстар анықталды',
            conflict_desc: 'Құжаттарда сәйкес емес деректер табылды:',
            metadata: 'Өңдеу статистикасы',
            files_processed: 'Талданған файлдар',
            final_assembly: 'Құжатты жинау',
            tokens: 'Өңделген токендер',
            duration: 'Өңдеу уақыты',
            ms: 'мс'
        }
    };
    
    const t = labels[currentLang] || labels.ru;

    const makeFact = (fact, fieldName = '') => {
        // Handle plain string (spec format fallback)
        if (typeof fact === 'string') {
            return fact && fact.length > 1 && !fact.includes('отсутству')
                ? `<span>${fact}</span>`
                : `<span style="color:var(--text-muted);font-style:italic;">${t.no_data}</span>`;
        }
        if (!fact) return `<span style="color:var(--text-muted);font-style:italic;">${t.no_data}</span>`;
        // Rich format: if text is empty but source is present, show "Источник: …" hint
        if (!fact.text || !fact.text.trim()) {
            if (fact.source && !fact.source.includes('файл.txt')) {
                return `<span style="color:var(--text-muted);font-style:italic;">Нет текста</span><span class="source-badge">📄 ${fact.source}</span>`;
            }
            return `<span style="color:var(--text-muted);font-style:italic;">${t.no_data}</span>`;
        }
        const badgeStyle = (fact.source === 'Нет источника' || (fact.source || '').includes('отсутству'))
            ? 'background:#F3F4F6;color:#6B7280;' : '';
        let html = `<span>${fact.text}</span><span class="source-badge" style="${badgeStyle}">📄 ${fact.source || '—'}</span>`;
        if (fact.has_conflict) {
            const conflictId = `conflict-${fieldName}-${Math.random().toString(36).substr(2, 9)}`;
            // Parse conflict details to extract sources
            const sources = parseConflictDetails(fact.conflict_details);
            html += `
            <div style="margin-top:8px;margin-bottom:8px;background:var(--conflict-bg);border-left:4px solid #F59E0B;padding:10px 14px;font-size:13.5px;color:var(--conflict-text);border-radius:0 4px 4px 0;cursor:pointer;"
                 onclick='showConflictFromData(${JSON.stringify(fact.conflict_details).replace(/'/g, "&#39;")})'>
                <strong style="display:block;margin-bottom:4px;">⚠️ ${t.conflict_click}:</strong>
                ${fact.conflict_details}
            </div>`;
        }
        return html;
    };

    // Helper to parse conflict details
    function parseConflictDetails(details) {
        const sources = [];
        const regex = /([^;]+?) — ([^;]+)/g;
        let match;
        while ((match = regex.exec(details)) !== null) {
            const fileLine = match[1].trim();
            const value = match[2].trim();
            const fileMatch = fileLine.match(/(.+?),\s*\[(\d+)\]/);
            if (fileMatch) {
                sources.push({ file: fileMatch[1], line: fileMatch[2], value });
            } else {
                sources.push({ file: fileLine, line: '?', value });
            }
        }
        return sources.length > 0 ? sources : [{ file: 'Неизвестно', line: '?', value: details }];
    }

    window.showConflictFromData = function(conflictDetails) {
        const sources = parseConflictDetails(conflictDetails);
        const modal = document.getElementById('conflict-modal');
        const container = document.getElementById('diff-container');

        const diffHtml = sources.map((src, idx) => `
            <div class="diff-box ${idx === 0 ? 'highlight' : ''}">
                <h4>${src.file}, строка ${src.line}</h4>
                <p>${src.value}</p>
            </div>
        `).join('');

        container.innerHTML = diffHtml;
        modal.style.display = 'flex';
    };

    const makeFactList = (items) => {
        if (!items || items.length === 0) return `<p style="color:var(--text-muted);font-style:italic;">${t.no_data}</p>`;
        // Filter out items that are completely empty (empty string or object with empty text)
        const nonEmpty = items.filter(item => {
            if (typeof item === 'string') return item.trim().length > 0;
            if (typeof item === 'object' && item !== null) {
                // Keep items that have real text OR a real conflict to show
                const hasText = (item.text || '').trim().length > 0;
                const hasSource = item.source && !item.source.includes('файл.txt');
                return hasText || hasSource || item.has_conflict;
            }
            return false;
        });
        if (nonEmpty.length === 0) return `<p style="color:var(--text-muted);font-style:italic;">${t.no_data}</p>`;
        return `<ul>${nonEmpty.map(item => `<li>${makeFact(item)}</li>`).join('')}</ul>`;
    };

    // Conflict summary banner
    const conflictsCount = meta.conflicts_found || 0;
    let conflictBannerHtml = '';
    if (conflictsCount > 0) {
        const fieldLabels = {
            timeline: t.timeline, budget: t.budget,
            technical_solution: t.technical_solution, architecture: t.architecture,
            goals: t.goals, requirements: t.requirements, team: t.team, risks: t.risks
        };
        const conflictItems = [];
        ['timeline','budget','technical_solution','architecture'].forEach(key => {
            const f = doc[key];
            if (f && f.has_conflict && f.conflict_details)
                conflictItems.push(`<li><strong>${fieldLabels[key]}:</strong> ${f.conflict_details}</li>`);
        });
        ['goals','requirements','team','risks'].forEach(key => {
            (doc[key] || []).forEach((f, i) => {
                if (f && f.has_conflict && f.conflict_details)
                    conflictItems.push(`<li><strong>${fieldLabels[key]} [${i+1}]:</strong> ${f.conflict_details}</li>`);
            });
        });
        conflictBannerHtml = `
        <div style="background:#FEF3C7;border:2px solid #F59E0B;border-radius:8px;padding:16px 20px;margin-bottom:24px;">
            <div style="display:flex;align-items:center;gap:10px;margin-bottom:4px;">
                <span style="font-size:22px;">⚠️</span>
                <strong style="font-size:16px;color:#92400E;">${t.conflicts_found}: ${conflictsCount}</strong>
            </div>
            <p style="margin:4px 0 0;color:#B45309;font-size:14px;">${t.conflict_desc}</p>
            <ul style="margin:8px 0 0;padding-left:20px;color:#92400E;font-size:13.5px;">${conflictItems.join('')}</ul>
        </div>`;
    }

    container.innerHTML = conflictBannerHtml + `
        <div class="result-section fade-in"><h3>${t.project_overview}</h3><p>${doc.project_overview || t.no_data}</p></div>
        <div class="result-section fade-in"><h3>${t.goals}</h3>${makeFactList(doc.goals)}</div>
        <div class="result-section fade-in"><h3>${t.requirements}</h3>${makeFactList(doc.requirements)}</div>
        <div class="result-section fade-in"><h3>${t.technical_solution}</h3><p>${makeFact(doc.technical_solution, 'technical_solution')}</p></div>
        <div class="result-section fade-in"><h3>${t.architecture}</h3><p>${makeFact(doc.architecture, 'architecture')}</p></div>
        <div class="result-section fade-in"><h3>${t.team}</h3>${makeFactList(doc.team)}</div>
        <div class="result-section fade-in">
            <h3>${t.timeline} & ${t.budget}</h3>
            <p><strong>${t.timeline}:</strong> ${makeFact(doc.timeline, 'timeline')}</p>
            <p><strong>${t.budget}:</strong> ${makeFact(doc.budget, 'budget')}</p>
        </div>
        <div class="result-section fade-in"><h3>${t.risks}</h3>${makeFactList(doc.risks)}</div>
    `;

    document.getElementById('meta-content').innerHTML = `
        <strong>${t.metadata}:</strong><br>
        ${t.files_processed}: ${(meta.llm_calls || 1) - 1}<br>
        ${t.final_assembly}: 1<br>
        ${t.conflicts_found}: <strong>${conflictsCount}</strong><br>
        ${t.tokens}: ${meta.total_tokens || 0} &nbsp;|&nbsp; ${t.duration}: ${meta.duration_ms || 0} ${t.ms}
    `;
}

// ─────────────────────────────────────────────────────────────────────────────
// BUTTONS
// ─────────────────────────────────────────────────────────────────────────────
document.getElementById('download-btn').addEventListener('click', () => {
    const blob = new Blob([JSON.stringify(currentRawData, null, 4)], { type: 'application/json' });
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url; a.download = 'map_reduce_doc.json'; a.click();
    URL.revokeObjectURL(url);
});

document.getElementById('load-another-btn').addEventListener('click', () => {
    showScreen('upload');
    switchTab('load');
});

document.getElementById('reset-btn').addEventListener('click', () => {
    document.getElementById('file-input').value = '';
    document.getElementById('json-file-input').value = '';
    showScreen('upload');
    switchTab('generate');
});
</script>
</body>
</html>
"""
    return HTMLResponse(content=html_content)

# ── Generate endpoint ─────────────────────────────────────────────
@app.post("/generate_document")
async def generate_document(files: List[UploadFile] = File(...), language: str = Form("ru")):
    files_data = []
    for file in files:
        content = await file.read()
        text = extract_text(file.filename, content)
        files_data.append({"filename": file.filename, "content": text})
    try:
        result = current_llm.generate_document(files_data, language=language)
        return result
    except Exception as e:
        error_msg = str(e)
        if "Connection error" in error_msg or "ConnectError" in error_msg:
            detail_msg = f"Не удалось подключиться к локальному серверу LM Studio по адресу {LOCAL_PC_IP}."
        else:
            detail_msg = f"Ошибка обработки: {error_msg}"
        raise HTTPException(status_code=500, detail=detail_msg)
