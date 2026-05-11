import io
import re
import ast
import json

# ─────────────────────────────────────────────────────────────────────────────
# FILE PARSERS (V2 Change 1): extract real signal from binary / code files
# before feeding them to the LLM. Cuts token cost dramatically on large
# projects (e.g. 11_omnicore_platform: ~190 KB → ~50-60 KB).
# ─────────────────────────────────────────────────────────────────────────────
def _parse_docx(content: bytes) -> str:
    """Extract paragraphs + table cell text from a .docx file."""
    try:
        from docx import Document
    except ImportError:
        return content.decode("utf-8", errors="ignore")
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        return f"[DOCX parse error: {e}]"
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _parse_pdf(content: bytes) -> str:
    """Extract text per page from a PDF."""
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # fallback
        except ImportError:
            return "[PDF parser not installed]"
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as e:
        return f"[PDF parse error: {e}]"
    out = []
    for i, page in enumerate(reader.pages, 1):
        try:
            t = (page.extract_text() or "").strip()
        except Exception:
            t = ""
        if t:
            out.append(f"--- page {i} ---\n{t}")
    return "\n\n".join(out)


def _parse_python(source: str) -> str:
    """
    AST-based extraction: module docstring, top-level constants,
    class/function signatures with their docstrings. Skips function bodies.
    Falls back to raw comments if AST fails.
    """
    out = []
    try:
        tree = ast.parse(source)
    except Exception:
        # fallback: keep only comments and docstrings-ish content
        keep = [ln for ln in source.splitlines()
                if ln.strip().startswith("#") or '"""' in ln or "'''" in ln]
        return "\n".join(keep)

    mod_doc = ast.get_docstring(tree)
    if mod_doc:
        out.append(f'"""{mod_doc}"""')

    for node in tree.body:
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            try:
                out.append(ast.unparse(node))
            except Exception:
                pass
        elif isinstance(node, ast.Assign):
            # top-level constants (UPPER_CASE)
            for tgt in node.targets:
                if isinstance(tgt, ast.Name) and tgt.id.isupper():
                    try:
                        out.append(ast.unparse(node))
                    except Exception:
                        pass
                    break
        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            try:
                args = ast.unparse(node.args)
            except Exception:
                args = "..."
            out.append(f"def {node.name}({args}):")
            d = ast.get_docstring(node)
            if d:
                out.append(f'    """{d}"""')
        elif isinstance(node, ast.ClassDef):
            bases = ""
            try:
                bases = ", ".join(ast.unparse(b) for b in node.bases)
            except Exception:
                pass
            out.append(f"class {node.name}({bases}):")
            d = ast.get_docstring(node)
            if d:
                out.append(f'    """{d}"""')
            for sub in node.body:
                if isinstance(sub, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    try:
                        args = ast.unparse(sub.args)
                    except Exception:
                        args = "..."
                    out.append(f"    def {sub.name}({args}):")
                    sd = ast.get_docstring(sub)
                    if sd:
                        out.append(f'        """{sd}"""')
    return "\n".join(out)


_JAVA_SIG_RE = re.compile(
    r"^\s*(?:public|private|protected|static|final|abstract|synchronized|\s)*"
    r"[\w<>\[\],\s]+\s+\w+\s*\([^)]*\)\s*(?:throws[\w\s,]+)?\s*[{;]",
    re.MULTILINE,
)
_JAVA_CLASS_RE = re.compile(
    r"^\s*(?:public|private|protected|abstract|final|\s)*"
    r"(?:class|interface|enum)\s+\w+[^{]*\{",
    re.MULTILINE,
)


def _parse_java(source: str) -> str:
    """
    Regex-based extraction: package, imports, javadoc/comments,
    class/interface headers, method signatures. Skips bodies.
    """
    out = []
    # package + imports
    for m in re.finditer(r"^\s*(package|import)\s+[^;]+;", source, re.MULTILINE):
        out.append(m.group(0).strip())

    # javadoc /** ... */ blocks
    for m in re.finditer(r"/\*\*.*?\*/", source, re.DOTALL):
        out.append(m.group(0))

    # // comments that look informative (skip trivial 1-word ones)
    for m in re.finditer(r"//[^\n]{6,}", source):
        out.append(m.group(0).strip())

    # class / interface / enum headers (without body)
    for m in _JAVA_CLASS_RE.finditer(source):
        out.append(m.group(0).rstrip("{").strip())

    # method signatures (drop the opening brace/semicolon)
    for m in _JAVA_SIG_RE.finditer(source):
        sig = m.group(0).rstrip("{;").strip()
        # filter noise: skip if it's actually an `if (...)` or similar control
        if re.match(r"^\s*(if|for|while|switch|catch|return)\b", sig):
            continue
        out.append(sig)

    return "\n".join(out)


def _flatten_json(content: bytes) -> str:
    """
    Convert a JSON file into flat "key: value" lines so the MAP model
    receives simple strings instead of nested dicts/arrays.
    Example: {"tech_stack": ["Python", "FastAPI"]} → "tech_stack: Python\ntech_stack: FastAPI"
    """
    try:
        data = json.loads(content.decode("utf-8", errors="ignore"))
    except Exception:
        return content.decode("utf-8", errors="ignore")

    lines: list = []

    def _walk(obj, prefix: str = ""):
        if isinstance(obj, dict):
            for k, v in obj.items():
                new_prefix = f"{k}: " if not prefix else f"{prefix}{k}: "
                _walk(v, new_prefix)
        elif isinstance(obj, list):
            for item in obj:
                _walk(item, prefix)
        else:
            val = str(obj).strip()
            if val:
                lines.append(f"{prefix}{val}" if prefix else val)

    _walk(data)
    return "\n".join(lines)


def _parse_markdown(text: str) -> str:
    """
    Strip heavy non-prose content from Markdown files so the MAP model
    receives readable text instead of SQL, protobuf, ASCII art, and YAML.

    Keeps: section headers (##), bullet points, regular paragraphs.
    Removes:
      • Fenced code blocks  (``` ... ```)
      • Lines that are mostly box-drawing characters (ASCII art diagrams)
      • Markdown table divider rows (|---|---|)
      • Trailing whitespace / excessive blank lines
    """
    # 1. Remove fenced code blocks (``` ... ``` or ~~~ ... ~~~)
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"~~~[\s\S]*?~~~", "", text)

    lines_out = []
    prev_blank = False
    for line in text.splitlines():
        stripped = line.strip()

        # Skip table divider rows: |---|---|
        if re.match(r"^\|[\s\-\|:]+\|$", stripped):
            continue

        # Skip lines that are ≥ 40% box-drawing / ASCII-art characters
        art_chars = sum(1 for c in stripped if c in "│├└─┐┘┌┤┬┴┼╔╗╚╝║═╠╣╦╩╬▼▲►◄")
        if stripped and art_chars / len(stripped) >= 0.4:
            continue

        # Collapse multiple blank lines into one
        if stripped == "":
            if not prev_blank:
                lines_out.append("")
            prev_blank = True
        else:
            lines_out.append(line.rstrip())
            prev_blank = False

    return "\n".join(lines_out).strip()


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch a file to the right parser based on extension."""
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            return _parse_docx(content)
        if name.endswith(".pdf"):
            return _parse_pdf(content)
        # textual formats from here on
        text = content.decode("utf-8", errors="ignore")
        if name.endswith(".py"):
            # Prefix tells the MAP model this is source code, not a spec document
            return "[SOURCE CODE — extract only technical_solution/architecture]\n" + _parse_python(text)
        if name.endswith(".java"):
            return "[SOURCE CODE — extract only technical_solution/architecture]\n" + _parse_java(text)
        if name.endswith(".json"):
            # Flatten nested JSON to "key: value" lines — prevents MAP from
            # copying complex objects verbatim as fact values
            return _flatten_json(content)
        if name.endswith(".md"):
            # Strip code blocks and ASCII art — large .md files like architecture.md
            # contain massive SQL/protobuf/YAML sections that overwhelm the model
            return _parse_markdown(text)
        # .txt and anything else → raw text
        return text
    except Exception as e:
        return f"[parser error for {filename}: {e}]"

